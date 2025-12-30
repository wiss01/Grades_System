# updates: missing forgot_pw  added successfully
# pip install opencv-python

# pip install deepface

# if not working use : pip install python-deepface

import csv
import cv2  # do imgs prccessing affiliated to deepface
from deepface import DeepFace


from tkinter import *
# interract with the file dialog os for exmp: import,export,os.remove(path),os.name ...
import os
import csv
import time
import tkinter as tk
# affiliated or relatedto windows os filedialog##
from tkinter import ttk, filedialog

# allow us to: open an ask window (for puting name... )
from tkinter import messagebox, simpledialog


from tkinter import Tk, Label

from tkinter import PhotoImage  # to just import the imaging Module from tkinter
from PIL import Image  # just to be able to import the image itself ##
# just to be able to manage the img (open, save, modify)
from PIL import ImageTk


from tkcalendar import DateEntry

import tkinter.font as tkFont  # for text fonts or style or size

from PIL import ImageSequence  # for the Gif


from tkinter.scrolledtext import ScrolledText
import tkinter as tk
from tkinter import ttk, filedialog, simpledialog, messagebox, scrolledtext
from datetime import datetime
import re
import os
import csv
from tkinter import Entry
import tkinter.messagebox as mb
import webbrowser
from tkinter import messagebox


import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.shared import Pt
from builtins import Exception


import os
import itertools
import random
import csv
from tkinter import Tk, Label, Entry, Button, messagebox
from tkinter import ttk
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.shared import Pt

import tkinter as tk
from tkinter import ttk
import os
from tkinter.filedialog import askdirectory

import subprocess

# "lib cut"


CSV_FILE = "log_data1.csv"  # to load the csv file


def forgot_password():

    # class belongs to opencv to be able to create a capturing video frame,
    cap = cv2.VideoCapture(0)
    # and (0) arg is the index of the cam (0 is the 1st cam , 1 is the 2nd cam ... )

    cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
    # the width and the height of the video capture frame or window
    cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)

    counter = 0
    face_match = False

    # Read the photo database from CSV file
    log_data = []
    with open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\log_data1.csv', 'r') as file:
        # creating a csv reader or Cursor (read and extract data from each iterated row in "file")
        reader = csv.reader(file)
        for row in reader:
            # append : add an ele to the end of "log_data = []" list
            log_data.append(row[4])

    while True:  # infinite loop looping until a break statment comes
        # frame: read the video frame, ret: true if the frame was readed successfully
        ret, frame = cap.read()

        if ret:  # if ture
            # every 30 (60 or 90) capture the facial recognition check is trigered
            if counter % 30 == 0:
                for photo_path in log_data:
                    reference_img = cv2.imread(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\face_photo\reference.jpg')
                    # assigned any img in the list as a ref img to start with

                    try:
                        if DeepFace.verify(frame, reference_img.copy())['verified']:

                            # The ['verified'] syntax is used to access the value associated with the 'verified' key in the result of the DeepFace.verify() function call.

                            face_match = True

                            break
                        else:
                            face_match = False
                    except:
                        face_match = False

            counter += 1  # if false then pass to the next frame and after 30 captures the check proc trigered again

            if face_match:

                print("your face_photo is match !")

                cv2.putText(frame, "Login successfully", (20, 450),cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 255, 0), 3)

                # delay=5000

                root.destroy()
                main1()

            else:
                cv2.putText(frame, "Access Denied", (20, 450),cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 0, 255), 3)

            cv2.imshow("video", frame)  # the frame title is 'video'

        # exe (the break ) after 1ms from the key pressing ('q')
        key = cv2.waitKey(1)
        if key == ord("q"):  # if 'q'
            break

    cv2.destroyAllWindows()  # close all opencv windows


def check_credentials(username_entry, password_entry, tree):
    username = username_entry.get()
    password = password_entry.get()
    is_valid = False

    for item in tree.get_children():
        data = tree.item(item)["values"]  # issue
        if data[0] == username and data[1] == password:
            is_valid = True
            break

    if is_valid:
        messagebox.showinfo("Login Successful", "Welcome to the application!")
        root.destroy()

        # Here put your app source code (start)

        print("This is the Login step!, Login button is clicked!, Credentials checked successfully!")

        main1()

        # Here put your app source code (end)

    else:
        messagebox.showerror("Login Failed", "Incorrect username or password.")


def check_admin_credentials(admin_password_entry, tree):
    password = admin_password_entry.get()

    if password == "admin":
        print("Admin is logged !")
        admin_login_window.destroy()
        logs_manager(tree)
    else:
        messagebox.showerror("Admin Login Failed", "Incorrect password.")


def open_admin_login_window(tree):
    global admin_login_window
    admin_login_window = tk.Toplevel(root)
    admin_login_window.title("Admin Login")

    admin_password_label = tk.Label(admin_login_window, text="Password:")
    admin_password_label.grid(row=0, column=0, padx=10, pady=10, sticky="e")
    admin_password_entry = tk.Entry(admin_login_window, show="*")
    admin_password_entry.grid(row=0, column=1, padx=10, pady=10)

    admin_login_button = tk.Button(admin_login_window, text="Login",      command=lambda: check_admin_credentials(admin_password_entry, tree))
    # columnspan=2: this means that the button will appears inthe middle of the entry above
    admin_login_button.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    admin_back_button = tk.Button(admin_login_window, text="Back", command=admin_login_window.destroy)
    admin_back_button.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

    admin_login_window.lift()


def load_from_csv(tree):
    try:
        # 'r': open it in the read mode ; newline="" : to handel newline syntax in diff OS
        with open(CSV_FILE, "r", newline="") as f:
            reader = csv.reader(f)  # create a reader or a cursor
            for row in reader:  # attach the cursor theads to each row in f
                # "": the 1st item will be in the root index position ,"end":the next item will be just at the side (means there is no espaces), "values=row": means, take the value from the current row (then put it into the current item)
                tree.insert("", "end", values=row)
    except FileNotFoundError:
        pass


def save_data(entry_fields, tree):
    new_data = [entry.get() for entry in entry_fields]

    if "" in new_data:
        messagebox.showerror("Save Error", "All fields must be filled.")
        return

    tree.insert("", "end", values=new_data)

    # Save data to the CSV file
    with open(CSV_FILE, "a", newline="") as f:
        writer = csv.writer(f)
        # This is a method of the CSV writer object that writes a single row of data to the CSV file.
        writer.writerow(new_data)

    messagebox.showinfo("Save Successful", "Data saved successfully.")


def update_csv(tree):
    data = [tree.item(item)["values"] for item in tree.get_children()]  # issue

    with open(CSV_FILE, "w", newline="") as f:
        writer = csv.writer(f)
        # This is a method of the CSV writer object that writes multi rows of data to the CSV file.
        writer.writerows(data)


def edit_data(entry_fields, tree):
    selected_item = tree.selection()
    if not selected_item:
        messagebox.showerror("Edit Error", "No item is selected.")
        return

    new_data = [entry.get() for entry in entry_fields]  # issue

    if "" in new_data:
        messagebox.showerror("Edit Error", "All fields must be filled.")
        return

    tree.item(selected_item, values=new_data)

    # Update data in the CSV file
    update_csv(tree)
    messagebox.showinfo("Edit Successful", "Data edited successfully.")


def delete_data(tree):
    selected_item = tree.selection()
    if not selected_item:
        messagebox.showerror("Delete Error", "No item is selected.")
        return

    tree.delete(selected_item)

    # Update data in the CSV file
    update_csv(tree)
    messagebox.showinfo("Delete Successful", "Data deleted successfully.")


def reset_data(entry_fields):
    for entry in entry_fields:
        entry.delete(0, tk.END)


def logs_manager(tree):
    logs_window = tk.Toplevel(root)
    logs_window.title("Logs Manager")

    # Entry fields
    field_names = ["Name/Username", "Password","Allowed_Section","Phone"]
    entry_fields = []

    for index, field_name in enumerate(field_names):  
        label = tk.Label(logs_window, text=f"{field_name}:") 
        label.grid(row=index, column=0, padx=10, pady=10, sticky="e")
        entry = tk.Entry(logs_window) 
        entry.grid(row=index, column=1, padx=10, pady=10)
        entry_fields.append(entry)

    # Treeview
    logs_tree = ttk.Treeview(logs_window, columns=field_names+["img_path"], show="headings")
    for name in field_names + ["img_path"]:
        logs_tree.heading(name, text=name)
    logs_tree.grid(row=4, column=0, columnspan=4,
                   padx=10, pady=10, sticky="news")

    # Load data from CSV into Treeview
    for item in tree.get_children():
        logs_tree.insert("", "end", values=tree.item(item)["values"])

    # Upload Image function
    def upload_image():
        selected_item = logs_tree.focus()
        if selected_item:
            image_path = filedialog.askopenfilename(title="Select an Image", filetypes=[
                                                    ("Image Files", "*.png *.jpg *.jpeg")])
            if image_path:
                # Update the img_path value in the selected item
                logs_tree.set(selected_item, "img_path", image_path)

                # Update the img_path value in the CSV file
                update_csv(logs_tree)

    # Buttons
    save_button = tk.Button(logs_window, text="Save",command=lambda: save_data(entry_fields, logs_tree))
    save_button.grid(row=5, column=0, padx=10, pady=10, sticky="news")

    edit_button = tk.Button(logs_window, text="Edit",command=lambda: edit_data(entry_fields, logs_tree))
    edit_button.grid(row=5, column=1, padx=10, pady=10, sticky="news")

    delete_button = tk.Button(logs_window, text="Delete", command=lambda: delete_data(logs_tree))
    delete_button.grid(row=5, column=2, padx=10, pady=10, sticky="news")

    reset_button = tk.Button(logs_window, text="Reset",command=lambda: reset_data(entry_fields))
    reset_button.grid(row=5, column=3, padx=10, pady=10, sticky="news")

    # Upload Image button
    upload_button = tk.Button(logs_window, text="Upload Image", command=upload_image)
    upload_button.grid(row=6, column=0, columnspan=4,
                       padx=10, pady=10, sticky="news")

# cut(end)


# Here put the source code of main1

def main1():

    # Define the actions for each button
    def action1():

        # Create the window
        window = tk.Toplevel(root)
        window.title('University Details')

        def uniText1():
            subWin1 = tk.Toplevel(window)
            subWin1.geometry('800x800')

            label = tk.Label(subWin1, text="The ISIMa (Institut Supérieur d'Informatique de Mahdia) is a higher education institution located in Mahdia, Tunisia. It was Established in 2004. It is also known as the Higher Institute of Computer Science of Mahdia . ISIMa is affiliated with the University of Monastir and offers programs and courses in the field of computer science and information technology. It provides education and training to students interested in pursuing careers in computer science, software engineering, and related fields. ISIMa is recognized for its academic programs and research activities in the field of computer science.")
            label.pack()
            print("This is About University consulting step!, About University button (def uniText1()) clicked!, Displaying Aubout University section successfully!")

        def uniText2():
            subWin2 = tk.Toplevel(window)
            subWin2.geometry('800x800')

            label = tk.Label(subWin2, text="uni internal system Description")
            label.pack()
            print("This is uni internal system consulting step!, About University button (def uniText2()) clicked!, Displaying uni internal system section successfully!")

        # Load the images
        from tkinter import Tk, Label, PhotoImage
        from PIL import Image, ImageTk

        image1 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\55.png'))
        image2 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\66.png'))

        # Create a style for the buttons
        style = ttk.Style()
        style.configure('Hover.TButton', background='gray90')

        # Create the buttons with associated actions and labels
        button1 = ttk.Button(window, image=image1, command=uniText1,text="uniText1Uni2", style='Hover.TButton', cursor='hand2')
        button1.grid(row=1, column=0, sticky='nsew')
        label1 = tk.Label(window, text="About University", width=30, height=2)
        label1.grid(row=2, column=0, sticky='nsew')

        button2 = ttk.Button(window, image=image2, command=uniText2,text="About Uni2", style='Hover.TButton', cursor='hand2')
        button2.grid(row=1, column=1, sticky='nsew')
        label2 = tk.Label(window, text="University internal System", width=30, height=2)
        label2.grid(row=2, column=1, sticky='nsew')

        # Run the main window event loop
        window.mainloop()

    def action2():

        print("You entered is the 'Student Management' section ! ")
        csv_file_name = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\students.csv'

        def read_csv():
            data = []
            with open(csv_file_name, mode='r', newline='', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile, delimiter=',')
                for row in reader:
                    data.append(row)
            return data

        def write_csv(data):
            with open(csv_file_name, mode='w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                for row in data:
                    writer.writerow(row)

        def save(edit_window, entries, treeview):
            new_data = [entry.get() for entry in entries]

            # check validation (start)
            pattern = r"^\d{4}-\d{2}-\d{2}$"
            date_of_birth = new_data[3]
            email_pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
            email = new_data[8]
            # Check if CIN contains  8 integers
            if not new_data[0].isdigit() or not len(new_data[0]) == 8:
                messagebox.showerror("Error", "CIN must contain  8 integers.")

            elif not (len(new_data[1]) >= 4 and len(new_data[1]) <= 10):
                messagebox.showerror("Error", "first name must contain at least 4 caracters and it should not pass 10 caracters.")

            elif not len(new_data[2]) >= 4 and len(new_data[2]) <= 10:
                messagebox.showerror("Error", "last name must contain at least 4 caracters and it should not pass 10 caracters.")

                # Regular expression to validate date in the format "YYYY-MM-DD"

            elif not re.match(pattern, date_of_birth):

                messagebox.showerror("Error", "Please respect this format:\n"
                                     "dd/mm/yy")

            elif not (len(new_data[4]) >= 4 and len(new_data[4]) <= 10):

                messagebox.showerror("Error", "Place of birth should be at least 4caracters and should not pass 10 caracters ")

            elif not (len(new_data[5]) >= 10 and len(new_data[5]) <= 30):

                messagebox.showerror("Error", "The Address should be at least 10caracters and should not pass 30 caracters ")

            elif not len(new_data[6]) <= 10:

                messagebox.showerror("Error", "class should not pass 10 caracters ")

            # Check if phone contains  8 integers
            elif not new_data[7].isdigit() or not len(new_data[7]) == 8:
                messagebox.showerror("Error", "phone number must contain  8 integers.")

            elif not re.match(email_pattern, email):

                messagebox.showerror("Error", "Please enter a valid email")

            # check validation (end)

            else:

                treeview.insert("", "end", values=new_data)
                edit_window.destroy()

                data = read_csv()
                data.append(new_data)
                write_csv(data)

                print("Data saved !")
                messagebox.showinfo("info", "Data saved !")

        def edit(edit_window, entries, treeview):
            item_id = treeview.selection()[0]
            new_data = [entry.get() for entry in entries]

            # check validation (start)
            pattern = r"^\d{4}-\d{2}-\d{2}$"
            date_of_birth = new_data[3]
            email_pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
            email = new_data[8]
            # Check if CIN contains  8 integers
            if not new_data[0].isdigit() or not len(new_data[0]) == 8:
                messagebox.showerror("Error", "CIN must contain  8 integers.")

            elif not (len(new_data[1]) >= 4 and len(new_data[1]) <= 10):
                messagebox.showerror("Error", "first name must contain at least 4 caracters and it should not pass 10 caracters.")

            elif not len(new_data[2]) >= 4 and len(new_data[2]) <= 10:
                messagebox.showerror("Error", "last name must contain at least 4 caracters and it should not pass 10 caracters.")

                # Regular expression to validate date in the format "YYYY-MM-DD"

            elif not re.match(pattern, date_of_birth):

                messagebox.showerror("Error", "Please respect this format:\n"
                                     "dd/mm/yy")

            elif not (len(new_data[4]) >= 4 and len(new_data[4]) <= 10):

                messagebox.showerror("Error", "Place of birth should be at least 4caracters and should not pass 10 caracters ")

            elif not (len(new_data[5]) >= 10 and len(new_data[5]) <= 30):

                messagebox.showerror("Error", "The Address should be at least 10caracters and should not pass 30 caracters ")

            elif not len(new_data[6]) <= 10:

                messagebox.showerror("Error", "class should not pass 10 caracters ")

            # Check if phone contains  8 integers
            elif not new_data[7].isdigit() or not len(new_data[7]) == 8:
                messagebox.showerror("Error", "phone number must contain  8 integers.")

            elif not re.match(email_pattern, email):

                messagebox.showerror("Error", "Please enter a valid email")

            # check validation (end)

            else:
                treeview.item(item_id, values=new_data)
                edit_window.destroy()

                data = read_csv()
                index = treeview.index(item_id)
                data[index] = new_data
                write_csv(data)

                print("Data edited !")
                messagebox.showinfo("info", "Data edited !")

        def get_index():

            selection = treeview.focus()
            index = treeview.index(selection)
            print(index)
            return index

        def delete(edit_window, entries, treeview):

            item_id = treeview.selection()[0]
            # item_id = treeview.focus()
            treeview.delete()
            # edit_window.destroy()

            index = get_index()
            treeview.delete(item_id)
            # data.pop(index) ## pop delete

            lines = []
            with open(csv_file_name, 'r') as file:
                reader = csv.reader(file)
                for i, row in enumerate(reader):
                    if i != index:
                        lines.append(row)

            with open(csv_file_name, 'w', newline='') as file:
                writer = csv.writer(file)
                writer.writerows(lines)

            print("Data deleted!")

            messagebox.showinfo("info", "Data deleted!")

        def reset(entries):
            for entry in entries:
                entry.delete(0, 'end')

            print("Data reseted!")
            ##messagebox.showinfo("info", "Data resetedv!")

        def on_double_click_wrapper(event):
            item_id = treeview.selection()[0]
            item = treeview.item(item_id)

            edit_window = tk.Toplevel(app)
            edit_window.title("Edit Student Details")

            topframe = tk.Frame(edit_window, padx=10, pady=10)
            topframe.grid(row=0, column=0, sticky='nsew')

            local_entries = []
            for i, text in enumerate(labels_text):
                label = ttk.Label(topframe, text=text)
                label.grid(column=0, row=i, padx=5, pady=5, sticky=tk.W)
                entry = ttk.Entry(topframe)
                entry.grid(column=1, row=i, padx=5, pady=5)
                entry.insert(0, item['values'][i])
                local_entries.append(entry)

            button_frame = tk.Frame(topframe, padx=5, pady=5)
            button_frame.grid(columnspan=2, sticky='ew')

            button_commands = [
                lambda: save(edit_window, local_entries, treeview),
                lambda: edit(edit_window, local_entries, treeview),
                lambda: delete(edit_window, local_entries, treeview),
                lambda: reset(local_entries)
            ]

            for i, (text, cmd) in enumerate(zip(button_texts, button_commands)):
                button = ttk.Button(button_frame, text=text, command=cmd)
                button.grid(column=i, row=0, padx=5, pady=5)

            topframe.columnconfigure(1, weight=1)

        # Search function
        def search():
            search_term = search_entry.get().lower()
            for item in treeview.get_children():  # item point on a row in the treeview each iteration
                treeview.delete(item)

            # row is a list , ' '.join(row).lower(): copy the current exsiting row '': between item an espace , we can't delete the first row cause this will be a sytax error
            filtered_data = [
                row for row in data if search_term in ' '.join(row).lower()]
            for row in filtered_data:
                treeview.insert("", "end", values=row)

            # config meth is called to update the text
            total_students_label.config(text=f"Total students: {len(filtered_data)}")

        # starting

        app = tk.Tk()
        app.title("Students List")

        students_list_title = tk.Label(app, text="Students List")
        students_list_title.grid(row=0, column=0, sticky='nsew')

        labels_text = ["CIN", "First ", "Last Name", "Date of Birth",
            "Place of Birth", "Address", "Class", "Phone", "Email"]
        button_texts = ["Save", "Edit", "Delete", "Reset"]
        entries = []

        treeview = ttk.Treeview(app, columns=labels_text,   show="headings", selectmode="browse")
        treeview.grid(row=1, column=0, sticky='nsew')

        for text in labels_text:
            treeview.heading(text, text=text)

        treeview.bind("<Double-1>", on_double_click_wrapper)

        data = read_csv()

        for row in data:
            treeview.insert("", "end", values=row)

        total_students_label = tk.Label(app, text=f"Total students: {len(data)}")
        total_students_label.grid(row=2, column=0, sticky='nsew')

        app.columnconfigure(0, weight=1)
        app.rowconfigure(0, weight=1)
        app.rowconfigure(1, weight=1)
        app.rowconfigure(2, weight=1)

        # search bar bloc

        # Create a frame for search bar
        search_frame = tk.Frame(app)
        search_frame.grid(row=3, column=0, sticky='nsew')

        search_label = ttk.Label(search_frame, text="Search:")
        search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

        search_entry = ttk.Entry(search_frame)
        search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')

        search_button = ttk.Button(search_frame, text="Search", command=search)
        search_button.grid(column=2, row=0, padx=5, pady=5)

        search_frame.columnconfigure(1, weight=1)

        # Configure the grid for the app
        app.columnconfigure(0, weight=1)
        app.rowconfigure(0, weight=1)
        app.rowconfigure(1, weight=1)
        app.rowconfigure(2, weight=1)
        app.rowconfigure(3, weight=1)

        app.mainloop()

    def action3():

        print("You entered 'Staff Management' section")

        csv_file_name = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\staff_members.csv'

        def read_csv():
            data = []
            with open(csv_file_name, mode='r', newline='', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile, delimiter=',')
                for row in reader:
                    data.append(row)
            return data

        def write_csv(data):
            with open(csv_file_name, mode='w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                for row in data:
                    writer.writerow(row)

        def save(edit_window, entries, treeview):
            new_data = [entry.get() for entry in entries]

            # check validation (start)

            email_pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
            email = new_data[4]
            # Check if CIN contains  8 integers
            if not new_data[0].isdigit() or not len(new_data[0]) == 8:
                messagebox.showerror("Error", "CIN must contain  8 integers.")

            elif not (len(new_data[1]) >= 5 and len(new_data[1]) <= 10):
                messagebox.showerror("Error", "first name must contain at least 5 caracters and it should not pass 10 caracters.")

            elif not (len(new_data[2]) >= 4 and len(new_data[2]) <= 10):
                messagebox.showerror("Error", "last name must contain at least 4 caracters and it should not pass 10 caracters.")

            # Check if phone contains  8 integers
            elif not new_data[3].isdigit() or not len(new_data[3]) == 8:
                messagebox.showerror("Error", "phone number must contain  8 integers.")

            elif not re.match(email_pattern, email):

                messagebox.showerror("Error", "Please enter a valid email")
                
            elif not (len(new_data[5]) >= 4 and len(new_data[5]) <= 10):
                messagebox.showerror("Error", "Ocuppation must contain at least 4 caracters and it should not pass 10 caracters.")

            # check validation (end)

            else:

                treeview.insert("", "end", values=new_data)
                edit_window.destroy()

                data = read_csv()
                data.append(new_data)
                write_csv(data)

                print("Data saved !")
                messagebox.showinfo("info", "Data saved !")

        def edit(edit_window, entries, treeview):
            item_id = treeview.selection()[0]
            new_data = [entry.get() for entry in entries]

            # check validation (start)

            email_pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
            email = new_data[4]
            # Check if CIN contains  8 integers
            if not new_data[0].isdigit() or not len(new_data[0]) == 8:
                messagebox.showerror("Error", "CIN must contain  8 integers.")

            elif not (len(new_data[1]) >= 5 and len(new_data[1]) <= 10):
                messagebox.showerror("Error", "first name must contain at least 5 caracters and it should not pass 10 caracters.")

            elif not (len(new_data[2]) >= 4 and len(new_data[1]) <= 10):
                messagebox.showerror("Error", "last name must contain at least 4 caracters and it should not pass 10 caracters.")

            # Check if phone contains  8 integers
            elif not new_data[3].isdigit() or not len(new_data[3]) == 8:
                messagebox.showerror("Error", "phone number must contain  8 integers.")

            elif not re.match(email_pattern, email):

                messagebox.showerror("Error", "Please enter a valid email")

            elif not (len(new_data[5]) >= 4 and len(new_data[5]) <= 10):
                messagebox.showerror("Error", "Ocuppation must contain at least 4 caracters and it should not pass 10 caracters.")

            # check validation (end)

            else:
                treeview.item(item_id, values=new_data)
                edit_window.destroy()

                data = read_csv()
                index = treeview.index(item_id)
                data[index] = new_data
                write_csv(data)

                print("Data edited !")
                messagebox.showinfo("info", "Data edited !")

        def get_index():

            selection = treeview.focus()
            index = treeview.index(selection)
            print(index)
            return index

        def delete(edit_window, entries, treeview):

            item_id = treeview.selection()[0]
            # item_id = treeview.focus()
            treeview.delete()
            # edit_window.destroy()

            index = get_index()
            treeview.delete(item_id)
            # data.pop(index) ## pop delete

            lines = []
            with open(csv_file_name, 'r') as file:
                reader = csv.reader(file)
                for i, row in enumerate(reader):
                    if i != index:
                        lines.append(row)

            with open(csv_file_name, 'w', newline='') as file:
                writer = csv.writer(file)
                writer.writerows(lines)

            print("Data deleted!")

            messagebox.showinfo("info", "Data deleted!")

        def reset(entries):
            for entry in entries:
                entry.delete(0, 'end')

            print("Data reseted !")
            # messagebox.showinfo("info","Data reseted !")

        def on_double_click_wrapper(event):
            # Directly get the id from the widget
            item_id = event.widget.selection()[0]
            item = event.widget.item(item_id)

            edit_window = tk.Toplevel(app)
            edit_window.title("Edit Staff Details")

            topframe = tk.Frame(edit_window, padx=10, pady=10)
            topframe.grid(row=0, column=0, sticky='nsew')

            local_entries = []
            for i, text in enumerate(labels_text):
                label = ttk.Label(topframe, text=text)
                label.grid(column=0, row=i, padx=5, pady=5, sticky=tk.W)
                entry = ttk.Entry(topframe)
                entry.grid(column=1, row=i, padx=5, pady=5)
                entry.insert(0, item['values'][i])
                local_entries.append(entry)

            button_frame = tk.Frame(topframe, padx=5, pady=5)
            button_frame.grid(columnspan=2, sticky='ew')

            button_commands = [
                lambda: save(edit_window, local_entries,event.widget),  # Pass the widget here
                lambda: edit(edit_window, local_entries,event.widget),  # And here
                lambda: delete(edit_window, local_entries,  event.widget),  # And here
                lambda: reset(local_entries)
            ]

            for i, (text, cmd) in enumerate(zip(button_texts, button_commands)):
                button = ttk.Button(button_frame, text=text, command=cmd)
                button.grid(column=i, row=0, padx=5, pady=5)

            topframe.columnconfigure(1, weight=1)

        # Search function

        def search():
            search_term = search_entry.get().lower()
            for item in treeview.get_children():
                treeview.delete(item)

            filtered_data = [
                row for row in data if search_term in ' '.join(row).lower()]
            for row in filtered_data:
                treeview.insert("", "end", values=row)

            total_staff_label.config(text=f"Total Staff Members : {len(filtered_data)}")

        # starting

        app = tk.Tk()
        app.title("Staff List")

        staff_list_title = tk.Label(app, text="Staff Members List")
        staff_list_title.grid(row=0, column=0, sticky='nsew')

        labels_text = ['CIN', 'First Name',
            'Last Name', 'Phone', 'Email', 'Occupation']
        button_texts = ["Save", "Edit", "Delete", "Reset"]
        entries = []

        treeview = ttk.Treeview(app, columns=labels_text,   show="headings", selectmode="browse")
        treeview.grid(row=1, column=0, sticky='nsew')

        for text in labels_text:
            treeview.heading(text, text=text)

        treeview.bind("<Double-1>", on_double_click_wrapper)

        data = read_csv()

        for row in data:
            treeview.insert("", "end", values=row)

        total_staff_label = tk.Label(app, text=f"Total Staff Members: {len(data)}")
        total_staff_label.grid(row=2, column=0, sticky='nsew')

        app.columnconfigure(0, weight=1)
        app.rowconfigure(0, weight=1)
        app.rowconfigure(1, weight=1)
        app.rowconfigure(2, weight=1)

        # search bar bloc

        # Create a frame for search bar
        search_frame = tk.Frame(app)
        search_frame.grid(row=3, column=0, sticky='nsew')

        search_label = ttk.Label(search_frame, text="Search:")
        search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

        search_entry = ttk.Entry(search_frame)
        search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')

        search_button = ttk.Button(search_frame, text="Search", command=search)
        search_button.grid(column=2, row=0, padx=5, pady=5)

        search_frame.columnconfigure(1, weight=1)

        # Configure the grid for the app
        app.columnconfigure(0, weight=1)
        app.rowconfigure(0, weight=1)
        app.rowconfigure(1, weight=1)
        app.rowconfigure(2, weight=1)
        app.rowconfigure(3, weight=1)

        app.mainloop()

    def action4():

        print("You entered 'University Resources Management' section")

        csv_file_name = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\uniResources.csv'

        def read_csv():
            data = []
            with open(csv_file_name, mode='r', newline='', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile, delimiter=',')
                for row in reader:
                    data.append(row)
            return data

        def write_csv(data):
            with open(csv_file_name, mode='w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                for row in data:
                    writer.writerow(row)

        def save(edit_window, entries, treeview):
            new_data = [entry.get() for entry in entries]

            
            # def is_float(value):
            # parts = value.split(".")
            # if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
            # return True
            # elif parts[0].isdigit():
            # return True
            # else:
            # return False

            # check validation (start)

            # check validation (start)
            ID_pattern = r'[A-Z]{3}-\d{3}$'
            Model_pattern = r'^[A-Za-z0-9_\-]+$'

            if not re.match(ID_pattern, new_data[0]):

                messagebox.showerror(  "Error",
                      "Invalid input. The input must match the pattern '^[A-Z]{3}-\d{3}$'.\n"
                      "The correct input should be three uppercase letters followed by a hyphen and three digits.",
                    )


            elif not re.match(Model_pattern, new_data[2]):

                messagebox.showerror(  "Error",
                          "Invalid input. The input must match the pattern '^[A-Za-z0-9_\-]+$'.\n"
                          "The correct input should start with a letter or underscore, followed by any number of letters, numbers, underscores, or hyphens.",
                        )

            elif not (len(new_data[1]) >= 5 and len(new_data[1]) <= 30):
                messagebox.showerror("Error", "The name of a university Resource must contain at least 5 caracters and it should not pass 30 caracters.")
            
            elif not (len(new_data[4]) >= 5 and len(new_data[4]) <= 15):
                messagebox.showerror("Error", "The Location Name must contain at least 5 caracters and it should not pass 15 caracters.")
            
            
            elif not (new_data[3].isdigit() and new_data[3] != ""):
                messagebox.showerror("Error", "Please enter a number !")
                
                
            elif not (new_data[5].isdigit() and new_data[5] != ""):
                messagebox.showerror("Error", "Please enter a number !")

            # check validation (end)

            else:
                treeview.insert("", "end", values=new_data)
                edit_window.destroy()

                data = read_csv()
                data.append(new_data)
                write_csv(data)

                print("Data saved !")
                messagebox.showinfo("info", "Data saved !")

        def edit(edit_window, entries, treeview):
            
            item_id = treeview.selection()[0]
            new_data = [entry.get() for entry in entries]
            value = new_data[3]
            # check validation (start)
            ID_pattern = r'[A-Z]{3}-\d{3}$'
            Model_pattern = r'^[A-Za-z0-9_\-]+$'

            if not re.match(ID_pattern, new_data[0]):

                messagebox.showerror(  "Error",
                      "Invalid input. The input must match the pattern '^[A-Z]{3}-\d{3}$'.\n"
                      "The correct input should be three uppercase letters followed by a hyphen and three digits.",
                    )


            elif not re.match(Model_pattern, new_data[2]):

                messagebox.showerror(  "Error",
                          "Invalid input. The input must match the pattern '^[A-Za-z0-9_\-]+$'.\n"
                          "The correct input should start with a letter or underscore, followed by any number of letters, numbers, underscores, or hyphens.",
                        )

            elif not (len(new_data[1]) >= 5 and len(new_data[1]) <= 30):
                messagebox.showerror("Error", "The name of a university Resource must contain at least 5 caracters and it should not pass 30 caracters.")
            
            elif not (len(new_data[4]) >= 5 and len(new_data[4]) <= 15):
                messagebox.showerror("Error", "The Location Name must contain at least 5 caracters and it should not pass 15 caracters.")
            
            
            elif not (value.isdigit() and value != ""):
                messagebox.showerror("Error", "Please enter a number !")
                
                
            elif not (new_data[5].isdigit() and value != ""):
                messagebox.showerror("Error", "Please enter a number !")

            # check validation (end)

            else:
                treeview.item(item_id, values=new_data)
                edit_window.destroy()
                data = read_csv()
                index = treeview.index(item_id)
                data[index] = new_data
                write_csv(data)
                print("Data edited !")
                messagebox.showinfo("info", "Data edited !")

        def get_index():

            selection = treeview.focus()
            index = treeview.index(selection)
            print(index)
            return index

        def delete(edit_window, entries, treeview):

            item_id = treeview.selection()[0]
            # item_id = treeview.focus()
            treeview.delete()
            # edit_window.destroy()

            index = get_index()
            treeview.delete(item_id)
            # data.pop(index) ## pop delete

            lines = []
            with open(csv_file_name, 'r') as file:
                reader = csv.reader(file)
                for i, row in enumerate(reader):
                    if i != index:
                        lines.append(row)

            with open(csv_file_name, 'w', newline='') as file:
                writer = csv.writer(file)
                writer.writerows(lines)

            print("Data deleted!")

            messagebox.showinfo("info", "Data deleted!")

        def reset(entries):
            for entry in entries:
                entry.delete(0, 'end')

        def on_double_click_wrapper(event):
            # Directly get the id from the widget
            item_id = event.widget.selection()[0]
            item = event.widget.item(item_id)

            edit_window = tk.Toplevel(app)
            edit_window.title("Edit Material Ressouce Details")

            topframe = tk.Frame(edit_window, padx=10, pady=10)
            topframe.grid(row=0, column=0, sticky='nsew')

            local_entries = []
            for i, text in enumerate(labels_text):
                label = ttk.Label(topframe, text=text)
                label.grid(column=0, row=i, padx=5, pady=5, sticky=tk.W)
                entry = ttk.Entry(topframe)
                entry.grid(column=1, row=i, padx=5, pady=5)
                entry.insert(0, item['values'][i])
                local_entries.append(entry)

            button_frame = tk.Frame(topframe, padx=5, pady=5)
            button_frame.grid(columnspan=2, sticky='ew')

            button_commands = [
                lambda: save(edit_window, local_entries,event.widget),  # Pass the widget here
                lambda: edit(edit_window, local_entries,event.widget),  # And here
                lambda: delete(edit_window, local_entries,  event.widget),  # And here
                lambda: reset(local_entries)
            ]

            for i, (text, cmd) in enumerate(zip(button_texts, button_commands)):
                button = ttk.Button(button_frame, text=text, command=cmd)
                button.grid(column=i, row=0, padx=5, pady=5)

            topframe.columnconfigure(1, weight=1)

        # Search function

        def search():
            search_term = search_entry.get().lower()
            for item in treeview.get_children():
                treeview.delete(item)

            filtered_data = [
                row for row in data if search_term in ' '.join(row).lower()]
            for row in filtered_data:
                treeview.insert("", "end", values=row)

            total_Ressouce_label.config(text=f"Total Ressouce Members : {len(filtered_data)}")

        # starting

        app = tk.Tk()
        app.title("University Material Resources List")

        Ressouce_list_title = tk.Label(app, text="Material Resources List")
        Ressouce_list_title.grid(row=0, column=0, sticky='nsew')

        labels_text = ["Resource ID",
                       "Resource Name",
                       "Resource Model",
                       "Resource Quantity",
                       "Resource Location",
                       "Cost(TND)"]
        button_texts = ["Save", "Edit", "Delete", "Reset"]
        entries = []

        treeview = ttk.Treeview(app, columns=labels_text,   show="headings", selectmode="browse")
        treeview.grid(row=1, column=0, sticky='nsew')

        for text in labels_text:
            treeview.heading(text, text=text)

        treeview.bind("<Double-1>", on_double_click_wrapper)

        data = read_csv()

        for row in data:
            treeview.insert("", "end", values=row)

        total_Ressouce_label = tk.Label(app, text=f"Total Material Resources: {len(data)}")
        total_Ressouce_label.grid(row=2, column=0, sticky='nsew')

        app.columnconfigure(0, weight=1)
        app.rowconfigure(0, weight=1)
        app.rowconfigure(1, weight=1)
        app.rowconfigure(2, weight=1)

        # search bar bloc

        # Create a frame for search bar
        search_frame = tk.Frame(app)
        search_frame.grid(row=3, column=0, sticky='nsew')

        search_label = ttk.Label(search_frame, text="Search:")
        search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

        search_entry = ttk.Entry(search_frame)
        search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')

        search_button = ttk.Button(search_frame, text="Search", command=search)
        search_button.grid(column=2, row=0, padx=5, pady=5)

        search_frame.columnconfigure(1, weight=1)

        # Configure the grid for the app
        app.columnconfigure(0, weight=1)
        app.rowconfigure(0, weight=1)
        app.rowconfigure(1, weight=1)
        app.rowconfigure(2, weight=1)
        app.rowconfigure(3, weight=1)

        app.mainloop()

    def action5():
        import tkinter as tk
        from tkinter import ttk
        from PIL import Image, ImageTk

        # Create the window
        window = tk.Toplevel(root)
        window.title('TimeTables Dashbord')

        def courses_rooms():
                print("You entered is the 'Courses & Rooms Management' section ! ")
                csv_file_name = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\Lectures.csv'
        
                def read_csv():
                    data = []
                    with open(csv_file_name, mode='r', newline='', encoding='utf-8') as csvfile:
                        reader = csv.reader(csvfile, delimiter=',')
                        for row in reader:
                            data.append(row)
                    return data
        
                def write_csv(data):
                    with open(csv_file_name, mode='w', newline='', encoding='utf-8') as csvfile:
                        writer = csv.writer(csvfile)
                        for row in data:
                            writer.writerow(row)
        
                def save(edit_window, entries, treeview):
                    new_data = [entry.get() for entry in entries]

                    ### check validation (start)
  
                    if not len(new_data[1]) >= 4 and len(new_data[1]) <= 10:
                        messagebox.showerror("Error", "Course Name must contain at least 4 caracters and it should not pass 10 caracters.")

                    elif not len(new_data[3]) >= 2 and len(new_data[3]) <= 10:
                        messagebox.showerror("Error", "Room name must contain at least 2 caracters and it should not pass 10 caracters.")
        
                    elif not (new_data[0].isdigit() and new_data[0] != ""):
                        messagebox.showerror("Error", "Please enter a number in the 1st entry !")
                        
                    elif not (new_data[2].isdigit() and new_data[2] != ""):
                        messagebox.showerror("Error", "Please enter a number in the 3rd entry !")    
                    
                    elif not (new_data[4].isdigit() and new_data[4] != ""):
                        messagebox.showerror("Error", "Please enter a number in the 5th entry !")   

                    ### check validation (end)
        
                    else:

                        treeview.insert("", "end", values=new_data)
                        edit_window.destroy()
        
                        data = read_csv()
                        data.append(new_data)
                        write_csv(data)
        
                        print("Data saved !")
                        messagebox.showinfo("info", "Data saved !")
        
                def edit(edit_window, entries, treeview):
                    item_id = treeview.selection()[0]
                    new_data = [entry.get() for entry in entries]
                    ### check validation (start)
  
                    if not len(new_data[1]) >= 4 and len(new_data[1]) <= 10:
                        messagebox.showerror("Error", "Course Name must contain at least 4 caracters and it should not pass 10 caracters.")

                    elif not len(new_data[3]) >= 2 and len(new_data[3]) <= 10:
                        messagebox.showerror("Error", "Room name must contain at least 2 caracters and it should not pass 10 caracters.")
        
                    elif not (new_data[0].isdigit() and new_data[0] != ""):
                        messagebox.showerror("Error", "Please enter a number in the 1st entry !")
                        
                    elif not (new_data[2].isdigit() and new_data[2] != ""):
                        messagebox.showerror("Error", "Please enter a number in the 3rd entry !")    
                    
                    elif not (new_data[4].isdigit() and new_data[4] != ""):
                        messagebox.showerror("Error", "Please enter a number in the 5th entry !")   

                    ### check validation (end)
                    else:
                        # It allows you to modify the values displayed in the columns of the specified item.
                        treeview.item(item_id, values=new_data)
                        edit_window.destroy()
        
                        data = read_csv()
                        # index(item_id): Calls the index() method on the Treeview widget, passing the item_id as an argument. The item_id represents the identifier of the selected item for which you want to retrieve the index.
                        index = treeview.index(item_id)
                        data[index] = new_data
                        write_csv(data)
        
                        print("Data edited!")
                        messagebox.showinfo("info", "Data edited!")
        
                def get_index():
        
                    selection = treeview.focus()
                    index = treeview.index(selection)
                    print(index)
                    return index
        
                def delete(edit_window, entries, treeview):
        
                    item_id = treeview.selection()[0]
                    # item_id = treeview.focus()
                    treeview.delete()
                    # edit_window.destroy()
        
                    index = get_index()
                    treeview.delete(item_id)
                    # data.pop(index) ## pop delete
        
                    lines = []
                    with open(csv_file_name, 'r') as file:
                        reader = csv.reader(file)
                        for i, row in enumerate(reader):
                            if i != index:
                                lines.append(row)
        
                    with open(csv_file_name, 'w', newline='') as file:
                        writer = csv.writer(file)
                        writer.writerows(lines)
        
                    print("Data deleted!")
        
                    messagebox.showinfo("info", "Data deleted!")
        
                def reset(entries):
                    for entry in entries:
                        entry.delete(0, 'end')
        
                    print("Data reseted!")
                    ##messagebox.showinfo("info", "Data resetedv!")
        
                def on_double_click_wrapper(event):
                    item_id = treeview.selection()[0]
                    item = treeview.item(item_id)
        
                    edit_window = tk.Toplevel(app)
                    edit_window.title("Edit Courses & Rooms Details")
        
                    topframe = tk.Frame(edit_window, padx=10, pady=10)
                    topframe.grid(row=0, column=0, sticky='nsew')
        
                    local_entries = []
                    for i, text in enumerate(labels_text):
                        label = ttk.Label(topframe, text=text)
                        label.grid(column=0, row=i, padx=5, pady=5, sticky=tk.W)
                        entry = ttk.Entry(topframe)
                        entry.grid(column=1, row=i, padx=5, pady=5)
                        entry.insert(0, item['values'][i])
                        local_entries.append(entry)
        
                    button_frame = tk.Frame(topframe, padx=5, pady=5)
                    button_frame.grid(columnspan=2, sticky='ew')
        
                    button_commands = [
                        lambda: save(edit_window, local_entries, treeview),
                        lambda: edit(edit_window, local_entries, treeview),
                        lambda: delete(edit_window, local_entries, treeview),
                        lambda: reset(local_entries)
                    ]
        
                    for i, (text, cmd) in enumerate(zip(button_texts, button_commands)):
                        button = ttk.Button(button_frame, text=text, command=cmd)
                        button.grid(column=i, row=0, padx=5, pady=5)
        
                    topframe.columnconfigure(1, weight=1)
        
                # Search function
                def search():
                    search_term = search_entry.get().lower()
                    for item in treeview.get_children():  # item point on a row in the treeview each iteration
                        treeview.delete(item)
        
                    # row is a list , ' '.join(row).lower(): copy the current exsiting row '': between item an espace , we can't delete the first row cause this will be a sytax error
                    filtered_data = [
                        row for row in data if search_term in ' '.join(row).lower()]
                    for row in filtered_data:
                        treeview.insert("", "end", values=row)
        
                    # config meth is called to update the text
                    total_students_label.config(text=f"Total courses: {len(filtered_data)}")
        
                # starting
        
                app = tk.Tk()
                app.title("Courses & Rooms Management")
        
                students_list_title = tk.Label(app, text="Courses & Rooms Management")
                students_list_title.grid(row=0, column=0, sticky='nsew')
        
                labels_text = ["Unit Code","Course Name","Duration","Room","Romm Capacity"]
                button_texts = ["Save", "Edit", "Delete", "Reset"]
                entries = []
        
                treeview = ttk.Treeview(app, columns=labels_text,           show="headings", selectmode="browse")
                treeview.grid(row=1, column=0, sticky='nsew')
        
                for text in labels_text:
                    treeview.heading(text, text=text)
        
                treeview.bind("<Double-1>", on_double_click_wrapper)
        
                data = read_csv()
        
                for row in data:
                    treeview.insert("", "end", values=row)
        
                total_students_label = tk.Label(app, text=f"Total courses: {len(data)}")
                total_students_label.grid(row=2, column=0, sticky='nsew')
        
                app.columnconfigure(0, weight=1)
                app.rowconfigure(0, weight=1)
                app.rowconfigure(1, weight=1)
                app.rowconfigure(2, weight=1)
        
                # search bar bloc
        
                # Create a frame for search bar
                search_frame = tk.Frame(app)
                search_frame.grid(row=3, column=0, sticky='nsew')
        
                search_label = ttk.Label(search_frame, text="Search:")
                search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)
        
                search_entry = ttk.Entry(search_frame)
                search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')
        
                search_button = ttk.Button(search_frame, text="Search", command=search)
                search_button.grid(column=2, row=0, padx=5, pady=5)
        
                search_frame.columnconfigure(1, weight=1)
        
                # Configure the grid for the app
                app.columnconfigure(0, weight=1)
                app.rowconfigure(0, weight=1)
                app.rowconfigure(1, weight=1)
                app.rowconfigure(2, weight=1)
                app.rowconfigure(3, weight=1)
        
                app.mainloop()

        def professors():

            print("You entered is the 'Professors Management' section ! ")
            csv_file_name = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\Profs.csv'
    
            def read_csv():
                data = []
                with open(csv_file_name, mode='r', newline='', encoding='utf-8') as csvfile:
                    reader = csv.reader(csvfile, delimiter=',')
                    for row in reader:
                        data.append(row)
                return data
    
            def write_csv(data):
                with open(csv_file_name, mode='w', newline='', encoding='utf-8') as csvfile:
                    writer = csv.writer(csvfile)
                    for row in data:
                        writer.writerow(row)
    
            def save(edit_window, entries, treeview):
                new_data = [entry.get() for entry in entries]
                ### check validation (start)

                if not len(new_data[1]) >= 4 and len(new_data[1]) <= 10:
                    messagebox.showerror("Error", "Professor Name must contain at least 4 caracters and it should not pass 10 caracters.")
                elif not (new_data[0].isdigit() and new_data[0] != ""):
                    messagebox.showerror("Error", "Please enter a number in the 1st entry !")
            
                elif not new_data[2].isdigit() or not len(new_data[2]) == 8:
                    messagebox.showerror("Error", "phone number must contain  8 integers.")
                
                elif not (new_data[3].isdigit() and new_data[3] != ""):
                    messagebox.showerror("Error", "Please enter a number in the 5th entry !")   
                ### check validation (end)
    
                else:
                    treeview.insert("", "end", values=new_data)
                    edit_window.destroy()
    
                    data = read_csv()
                    data.append(new_data)
                    write_csv(data)
    
                    print("Data saved !")
                    messagebox.showinfo("info", "Data saved !")
    
            def edit(edit_window, entries, treeview):
                item_id = treeview.selection()[0]
                new_data = [entry.get() for entry in entries]
                ### check validation (start)

                if not len(new_data[1]) >= 4 and len(new_data[1]) <= 10:
                    messagebox.showerror("Error", "Professor Name must contain at least 4 caracters and it should not pass 10 caracters.")
                elif not (new_data[0].isdigit() and new_data[0] != ""):
                    messagebox.showerror("Error", "Please enter a number in the 1st entry !")
            
                elif not new_data[2].isdigit() or not len(new_data[2]) == 8:
                    messagebox.showerror("Error", "phone number must contain  8 integers.")
                
                elif not (new_data[3].isdigit() and new_data[3] != ""):
                    messagebox.showerror("Error", "Please enter a number in the 5th entry !")   
                ### check validation (end)
                else:
                    # It allows you to modify the values displayed in the columns of the specified item.
                    treeview.item(item_id, values=new_data)
                    edit_window.destroy()
    
                    data = read_csv()
                    # index(item_id): Calls the index() method on the Treeview widget, passing the item_id as an argument. The item_id represents the identifier of the selected item for which you want to retrieve the index.
                    index = treeview.index(item_id)
                    data[index] = new_data
                    write_csv(data)
    
                    print("Data edited!")
                    messagebox.showinfo("info", "Data edited!")
    
            def get_index():
    
                selection = treeview.focus()
                index = treeview.index(selection)
                print(index)
                return index
    
            def delete(edit_window, entries, treeview):
    
                item_id = treeview.selection()[0]
                # item_id = treeview.focus()
                treeview.delete()
                # edit_window.destroy()
    
                index = get_index()
                treeview.delete(item_id)
                # data.pop(index) ## pop delete
    
                lines = []
                with open(csv_file_name, 'r') as file:
                    reader = csv.reader(file)
                    for i, row in enumerate(reader):
                        if i != index:
                            lines.append(row)
    
                with open(csv_file_name, 'w', newline='') as file:
                    writer = csv.writer(file)
                    writer.writerows(lines)
    
                print("Data deleted!")
    
                messagebox.showinfo("info", "Data deleted!")
    
            def reset(entries):
                for entry in entries:
                    entry.delete(0, 'end')
    
                print("Data reseted!")
                ##messagebox.showinfo("info", "Data resetedv!")
    
            def on_double_click_wrapper(event):
                item_id = treeview.selection()[0]
                item = treeview.item(item_id)
    
                edit_window = tk.Toplevel(app)
                edit_window.title("Edit Professors Details")
    
                topframe = tk.Frame(edit_window, padx=10, pady=10)
                topframe.grid(row=0, column=0, sticky='nsew')
    
                local_entries = []
                for i, text in enumerate(labels_text):
                    label = ttk.Label(topframe, text=text)
                    label.grid(column=0, row=i, padx=5, pady=5, sticky=tk.W)
                    entry = ttk.Entry(topframe)
                    entry.grid(column=1, row=i, padx=5, pady=5)
                    entry.insert(0, item['values'][i])
                    local_entries.append(entry)
    
                button_frame = tk.Frame(topframe, padx=5, pady=5)
                button_frame.grid(columnspan=2, sticky='ew')
    
                button_commands = [
                    lambda: save(edit_window, local_entries, treeview),
                    lambda: edit(edit_window, local_entries, treeview),
                    lambda: delete(edit_window, local_entries, treeview),
                    lambda: reset(local_entries)
                ]
    
                for i, (text, cmd) in enumerate(zip(button_texts, button_commands)):
                    button = ttk.Button(button_frame, text=text, command=cmd)
                    button.grid(column=i, row=0, padx=5, pady=5)
    
                topframe.columnconfigure(1, weight=1)
    
            # Search function
            # Search function
            def search1():
                search_term = search_entry.get().lower()
                for item in treeview.get_children():  # item point on a row in the treeview each iteration
                    treeview.delete(item)
    
                # row is a list , ' '.join(row).lower(): copy the current exsiting row '': between item an espace , we can't delete the first row cause this will be a sytax error
                filtered_data = [
                    row for row in data if search_term in ' '.join(row).lower()]
                for row in filtered_data:
                    treeview.insert("", "end", values=row)
    
                # config meth is called to update the text
                total_students_label.config(text=f"Total Professors: {len(filtered_data)}")
        
            # starting
    
            app = tk.Tk()
            app.title("Professors Management")
    
            students_list_title = tk.Label(app, text="Professors Management")
            students_list_title.grid(row=0, column=0, sticky='nsew')
    
            labels_text = ["Unit Code","Professor Name","Phone","ID"]
            button_texts = ["Save", "Edit", "Delete", "Reset"]
            entries = []
    
            treeview = ttk.Treeview(app, columns=labels_text,       show="headings", selectmode="browse")
            treeview.grid(row=1, column=0, sticky='nsew')
    
            for text in labels_text:
                treeview.heading(text, text=text)
    
            treeview.bind("<Double-1>", on_double_click_wrapper)
    
            data = read_csv()
    
            for row in data:
                treeview.insert("", "end", values=row)
    
            total_students_label = tk.Label(app, text=f"Total Professors: {len(data)}")
            total_students_label.grid(row=2, column=0, sticky='nsew')
    
            app.columnconfigure(0, weight=1)
            app.rowconfigure(0, weight=1)
            app.rowconfigure(1, weight=1)
            app.rowconfigure(2, weight=1)
    
            # search bar bloc
            
            
            
            ##            # Create a frame for search bar
            ##search_frame = tk.Frame(app)
            ##search_frame.grid(row=3, column=0, sticky='nsew')
    ##
            ##search_label = ttk.Label(search_frame, text="Search:")
            ##search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)
    ##
            ##search_entry = ttk.Entry(search_frame)
            ##search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')
    ##
            ##search_button = ttk.Button(search_frame, text="Search", command=search)
            ##search_button.grid(column=2, row=0, padx=5, pady=5)
    ##
            ##search_frame.columnconfigure(1, weight=1)
            
            

            
            # Create a frame for search bar
            search_frame = tk.Frame(app)
            search_frame.grid(row=3, column=0, sticky='nsew')
    
            search_label = ttk.Label(search_frame, text="Search:")
            search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)
    
            search_entry = ttk.Entry(search_frame)
            search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')
    
            search_button = ttk.Button(search_frame, text="Search", command=search1)
            search_button.grid(column=2, row=0, padx=5, pady=5)
    
            search_frame.columnconfigure(1, weight=1)
    
            # Configure the grid for the app
            app.columnconfigure(0, weight=1)
            app.rowconfigure(0, weight=1)
            app.rowconfigure(1, weight=1)
            app.rowconfigure(2, weight=1)
            app.rowconfigure(3, weight=1)
    
            app.mainloop()
        def a3():
            import tkinter as tk
            from tkinter import ttk
            import os
            from tkinter.filedialog import askdirectory

            import subprocess

            def style_word():
                try:
                    file_path = fr'C:\Users\Administrateur\OneDrive\Desktop\p_Root\P2_Planification_Examination_Library\1_timetables & exams_P\TT2\5_generateTT\funsForTheMainProgTT\output_folder2\outputfill1.docx'
                    subprocess.Popen(['start', 'winword', file_path], shell=True)
                except Exception as e:
                    print(f"An error occurred: {str(e)}")

            def load_data(folder_path):
                data = []
                for file_name in os.listdir(folder_path):
                    if file_name.endswith(".docx"):
                        file_path = os.path.join(folder_path, file_name)
                        data.append(file_path)
                return data

            def update_treeview(data):
                treev.delete(*treev.get_children())
                for i, file_path in enumerate(data, start=1):
                    file_name = os.path.basename(file_path)
                    treev.insert("", "end", iid=i, text=file_name,   values=(file_name, file_path))

            def delete_files_in_folder():

                folder_path = fr"C:\Users\Administrateur\OneDrive\Desktop\p_Root\P2_Planification_Examination_Library\1_timetables & exams_P\TT2\5_generateTT\funsForTheMainProgTT\output_folder2"
                confirmation = messagebox.askokcancel(    "Confirm Deletion",
                    f"We will start new TT generation, so we should delete all files in folder {folder_path} will be deleted. Are you sure you want to proceed?"
                )

                if not confirmation:
                    return

                try:
                    files = os.listdir(folder_path)
                    for file_name in files:
                        file_path = os.path.join(folder_path, file_name)
                        os.remove(file_path)
                    messagebox.showinfo(        "Files Deleted", f"All files in folder {folder_path} have been deleted.")
                except Exception as e:
                    messagebox.showerror(        "Error", f"An error occurred: {str(e)}")

            # TT generator

            def TT_generator():

                  # 0 step : deleting the exsiting TT , to generate new ones

                    delete_files_in_folder()

                    # 1st step: Data extract

                    import csv
                    from tkinter import Tk, Label, Entry, Button, messagebox

                    # Function to load CSV data
                    def load_csv(file_path):
                        data = []
                        with open(file_path, 'r') as csvfile:
                            reader = csv.reader(csvfile)
                            for row in reader:
                                data.append(row)
                        return data

                    # Load the CSV data for the first database (Lectures)
                    lectures_data = load_csv(        fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\Lectures.csv')

                    # Load the CSV data for the second database (Profs)
                    profs_data = load_csv(        fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\Profs.csv')

                    # Function to preprocess the data (if needed)
                    def preprocess_data(data):
                        # Remove any empty rows or perform additional preprocessing steps
                        data = [row for row in data if any(row)]
                        return data

                    lectures_data = preprocess_data(lectures_data)
                    profs_data = preprocess_data(profs_data)

                    # Function to extract relevant information and store in a vertical list
                    def extract_information():

                        path = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\output.csv'

                        # Clear the existing output file
                        with open(path, 'w') as csvfile:
                            csvfile.truncate()

                        # Iterate through each unit code in the lectures data
                        for lecture_row in lectures_data[0:]:
                            unit_code = lecture_row[0]
                            lecture_name = lecture_row[1].replace(                ",", " ")  # Replace commas with spaces
                            room_name = lecture_row[3].replace(                ",", " ")  # Replace commas with spaces

                            # Find the corresponding professor(s) in the profs data
                            prof_names = []
                            for prof_row in profs_data[0:]:
                                if prof_row[0] == unit_code:
                                    # Replace commas with spaces
                                    prof_names.append(                        prof_row[1].replace(",", " "))

                            # Remove the ',' and 'id' in the first column
                            unit_code = unit_code.split(',')[0]

                            # Append the extracted information to the output file (excluding the first column)
                            with open(path, 'a') as csvfile:
                                writer = csv.writer(csvfile)
                                writer.writerow(                    [lecture_name, room_name] + prof_names)

                        # Read the data from the output file and remove empty lines between rows
                        with open(path, 'r') as csvfile:
                            data = csvfile.read()
                        with open(path, 'w') as csvfile:
                            csvfile.write(                '\n'.join(line.strip() for line in data.strip().splitlines() if line.strip()))

                        # Display success message
                        print("Information Extracted",  "Information has been extracted successfully!")

                        # step2: convert the csv to txt , cause it's the only way to remove comas from csv

                    import csv
                    from tkinter import Tk, Button, messagebox

                    # Function to load CSV data
                    def load_csv(file_path):
                        data = []
                        with open(file_path, 'r') as csvfile:
                            reader = csv.reader(csvfile)
                            for row in reader:
                                data.append(row)
                        return data

                    # Function to preprocess the data (if needed)
                    def preprocess_data(data):
                        # Remove any empty rows or perform additional preprocessing steps
                        data = [row for row in data if any(row)]
                        return data

                    # Function to transform the data and save as TXT file
                    def transform_to_txt():
                        # Load the CSV data for the first database (Lectures)
                        lectures_data = load_csv(            fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\output.csv')
                        lectures_data = preprocess_data(lectures_data)

                        # Clear the existing output file
                        with open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\output.txt', 'w') as txtfile:
                            txtfile.truncate()

                        # Process each row of the CSV data and write to the TXT file
                        for row in lectures_data:
                            # Join the elements in the row with spaces (instead of commas)
                            transformed_row = ' '.join(row)

                            # Write the transformed row to the TXT file
                            with open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\output.txt', 'a') as txtfile:
                                txtfile.write(transformed_row + '\n')

                        # Display success message
                        print("Transformation Complete","CSV file transformed and saved as TXT file!")

                    # step3 : generate permutation (apply the changes)

                    import os
                    import itertools
                    import random

                    def generate_permutations():
                        data_file = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\output.txt'

                        data = []

                        # Read data from the file and populate the data list
                        with open(data_file, 'r') as file:
                            for line in file:
                                data.append(line.strip())

                        # Print the data to verify
                        print(data)

                        # Create a folder to store the text files
                        folder_path = fr"C:\Users\Administrateur\OneDrive\Desktop\p_Root\P2_Planification_Examination_Library\1_timetables & exams_P\TT2\5_generateTT\funsForTheMainProgTT\output_folder2"
                        os.makedirs(folder_path, exist_ok=True)

                        # Take the number of permutations as input
                        num_permutations = int(            input("Enter the number of permutations to generate: "))

                        # Store each permutation in a text file
                        for i in range(num_permutations):
                            # Shuffle the data list to create a random arrangement of lines
                            random.shuffle(data)

                            file_name = f"permutation_{i}.txt"
                            file_path = os.path.join(folder_path, file_name)

                            # Open the file and write each line with the row index
                            with open(file_path, "w") as f:
                                for row_index, line in enumerate(data, start=1):
                                    f.write(f"{line}\n")

                        print(            f"Number of possible permutations: {num_permutations}")
                        print("Permutations stored in text files.")

                    # step4: generate timetables (apply the changes)

                    import os
                    from docx import Document
                    from docx.enum.table import WD_ALIGN_VERTICAL
                    from docx.oxml import parse_xml
                    from docx.shared import Pt

                    # def generateTT():
                    # Define the nsdecls function

                    def nsdecls(prefix):
                        return 'xmlns:{0}="urn:schemas-microsoft-com:office:word"'.format(prefix)

                    # Function to generate timetable using the content from a text file
                    def generate_timetable(file_path, output_file):
                        # Create a new Word document
                        doc = Document()

                        # Add a paragraph at the beginning
                        doc.add_paragraph(            "The Higher Institute of Computer Science of Mahdia isima")

                        # Table data
                        table_data = [
                            ['Time', '08:00-10:00', '10:00-12:00','12:00-13:00', '13:00-15:00', '15:00-17:00'],['Monday', '', '', '', '', ''],['Tuesday', '', '', '', '', ''],['Wednesday', '', '', '', '', ''],['Thursday', '', '', '', '', ''],['Friday', '', '', '', '', ''],
                        ]

                        # Add the table to the document
                        table = doc.add_table(            rows=len(table_data), cols=len(table_data[0]))

                        # Set table properties
                        table.style = 'Table Grid'

                        # Set cell vertical alignment and font size
                        for row in table.rows:
                            for cell in row.cells:
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(12)

                        # Set table cell borders
                        for row in table.rows:
                            for cell in row.cells:
                                for child in cell._element:
                                    if child.tag.endswith('tcPr'):
                                        child.append(                            parse_xml(r'<w:shd {} w:fill="auto"/>'.format(nsdecls('w'))))

                        # Read content from the text file
                        with open(file_path, 'r') as file:
                            content = file.read().splitlines()

                    # Populate the table with data from the text file
                        index = 0
                        for i in range(1, len(table_data)):
                            for j in range(1, len(table_data[i])):
                                # Check if the cell is not '12:00-13:00' before populating it with content
                                if table_data[i][j] == '' and table_data[0][j] != '12:00-13:00':
                                    if index < len(content):
                                        table_data[i][j] = content[index]
                                        index += 1

                    # Populate the table with data (rest of the code remains
                        # Populate the table with data
                        for i, row in enumerate(table_data):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = cell

                        # Add a paragraph at the end
                        doc.add_paragraph("Director's signature:")

                        # Save the document
                        doc.save(output_file)

                    # the inputfolder is the outputfolder , end generateTT

                    extract_information()
                    transform_to_txt()
                    generate_permutations()

                    # Main function

                    output_folder = fr"C:\Users\Administrateur\OneDrive\Desktop\p_Root\P2_Planification_Examination_Library\1_timetables & exams_P\TT2\5_generateTT\funsForTheMainProgTT\output_folder2"
                    # Get the list of all text files in the output_folder
                    text_files = [f for f in os.listdir(        output_folder) if f.endswith(".txt")]
                    if not text_files:
                        print("No text files found in the output folder.")
                    else:
                        for i, text_file in enumerate(text_files, 1):
                            input_file = os.path.join(output_folder, text_file)
                            output_file = os.path.join(                output_folder, f"outputfill{i}.docx")
                            generate_timetable(input_file, output_file)
                        print("Timetables generated successfully.")

                        # this for auto refresh (when u generate new TT)
                        folder_path = fr'C:\Users\Administrateur\OneDrive\Desktop\p_Root\P2_Planification_Examination_Library\1_timetables & exams_P\TT2\5_generateTT\funsForTheMainProgTT\output_folder2'
                        if folder_path:
                            data = load_data(folder_path)
                            update_treeview(data)

            root = tk.Tk()
            root.geometry("500x400")  # Set the initial size of the main window

            # Create the Treeview widget
            treev = ttk.Treeview(root)
            # anchor=tk.CENTER ,  when u maximize the window the treev stay in the center
            treev.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

            # Bind double-click event to the treeview
            style_word_button = ttk.Button(root, command=style_word, text="let's style it")
            style_word_button.pack()

            gen_button = ttk.Button(root, command=TT_generator, text="Let's generate Timetables")
            gen_button.pack()

            # this for initialisation (what u view at the first time)
            folder_path = fr'C:\Users\Administrateur\OneDrive\Desktop\p_Root\P2_Planification_Examination_Library\1_timetables & exams_P\TT2\5_generateTT\funsForTheMainProgTT\output_folder2'
            if folder_path:
                data = load_data(folder_path)
                update_treeview(data)

            root.mainloop()

            
        def a4():
            # Specify the initial directory for the file dialog
            initial_path = fr"C:\Users\Administrateur\OneDrive\Desktop\p_Root\P2_Planification_Examination_Library\1_timetables & exams_P\TT2\5_generateTT\funsForTheMainProgTT\output_folder2"

            # Show the file dialog with the specified initial directory and all files (*.*)
            file_path = filedialog.askopenfilename(title="Select File", initialdir=initial_path, filetypes=[("All Files", "*.*")])

            # Check if the user selected a file
            if file_path:
                print("Selected file:", file_path)

        # Load the images
        image1 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\8_course icon.png'))
        image2 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\9_professor icon.png'))
        image3 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\10_TT icon.png'))
        image4 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\12_archive_icon.png'))

        # Create the buttons, assigning each one its action
        style = ttk.Style()
        style.configure('Hover.TButton', background='gray90')

        button1 = ttk.Button(window, image=image1, command=courses_rooms,style='Hover.TButton', cursor='hand2')
        button1.grid(row=1, column=0, sticky='nsew')
        label1 = tk.Label(window, text="Courses & Rooms")
        label1.grid(row=2, column=0, sticky='nsew')

        button2 = ttk.Button(window, image=image2, command=professors, style='Hover.TButton', cursor='hand2')
        button2.grid(row=1, column=1, sticky='nsew')
        label2 = tk.Label(window, text="Professors")
        label2.grid(row=2, column=1, sticky='nsew')

        button3 = ttk.Button(window, image=image3, command=a3,style='Hover.TButton', cursor='hand2')
        button3.grid(row=1, column=2, sticky='nsew')
        label3 = tk.Label(window, text="TimeTables Generation")
        label3.grid(row=2, column=2, sticky='nsew')

        button4 = ttk.Button(window, image=image4, command=a4,style='Hover.TButton', cursor='hand2')
        button4.grid(row=3, column=1, sticky='nsew')
        label4 = tk.Label(window, text="TimeTables Archive")
        label4.grid(row=4, column=1, sticky='nsew')

        # Configure grid to expand properly when resizing
        for i in range(3):  # number of columns
            window.columnconfigure(i, weight=1)

        for i in range(4):  # number of rows
            window.rowconfigure(i, weight=1)

        # Start the Tkinter event loop
        window.mainloop()

        # Start the Tkinter event loop
        # root.mainloop()

    def action6():
 

        def administrator_section(data):

                    print("You entered 'Grades Admin section'")

                    messagebox.showinfo("info", "name of the Node-Treeview = name of the Top node in the treeview\n"

                                        "to allow the system to get the node-treeview Node Dict name automatically")

                    # Note: the behavior of the code tells ttk.Treeview modul to use nodes icons and nodes build-in functions
                    class NodeDetailWindow(tk.Toplevel):
                        def __init__(self, master, node_name, node_type, node_path):
                            super().__init__(master)
                            self.title("Node Detail")

                            # Create labels and display node details
                            name_label = tk.Label(self, text="Node Name:")
                            name_label.pack()
                            name_value = tk.Label(self, text=node_name)
                            name_value.pack()

                            type_label = tk.Label(self, text="Node Type:")
                            type_label.pack()
                            type_value = tk.Label(self, text=node_type)
                            type_value.pack()

                            path_label = tk.Label(self, text="Node Path:")
                            path_label.pack()
                            path_value = tk.Label(self, text=node_path)
                            path_value.pack()

                    root = tk.Tk()
                    root.title("Treeview Syst")

                    tree = ttk.Treeview(root)
                    tree.pack(side="top", fill="both", expand=True)

                    labels_text = ["Wrong-Number", "TP", "DS", "Exam"]

                    treeview1 = ttk.Treeview(root, columns=labels_text, show="headings", selectmode="browse")

                    def write_data_to_node(node_name, data):
                        filename = f"{node_name}.csv"

                        # Write data to the specific node's CSV file
                        with open(filename, mode='a', newline='') as file:
                            writer = csv.writer(file)
                            writer.writerow(data)

                    def get_active_tree_name(trees, trees_names):
                            active_tree = None
                            for tree in trees:
                                if tree.focus():
                                    active_tree = tree
                                    break

                            if active_tree is None:
                                print("No active tree found.")
                                return None

                            return print(trees_names[active_tree])

                    # tree (is node-treeview ) (this is the only active treeview ,   "tree=t2")

                    # treeview1 == simple treeview inside

                    def myDBdash(event, treeview_window):

                        def validate_input(new_value):
                            if new_value == "":
                                return True
                            try:
                 
                                error_label.grid_forget()  # Hide the error message if input is valid
                                return True
                            except ValueError:
                                # Show the error message below the entry
                                error_label.grid(column=1, row=i+1, padx=5, pady=2)
                                return False

                        selected_item = treeview_window.focus()
                        if selected_item:
                            node_name = tree.item(selected_item, 'text')

                        selected_item = tree.selection()[0]
                        node_id = selected_item
                        csv_file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                        def read_csv():

                            selected_item = tree.focus()
                            node_name = tree.item(selected_item, 'text')
                            node_id = selected_item
                            csv_file = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                            data = []
                            with open(csv_file, mode='r') as csvfile:
                                reader = csv.reader(csvfile)
                                for row in reader:
                                    data.append(row)
                            return data

                        def write_csv(data):

                            selected_item = tree.focus()
                            node_name = tree.item(selected_item, 'text')
                            node_id = selected_item
                            csv_file = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                            with open(csv_file, mode='w', newline='', encoding='utf-8') as csvfile:
                                writer = csv.writer(csvfile)
                                for row in data:
                                    writer.writerow(row)

                        def save(edit_window, entries):
                            new_data = [entry.get() for entry in entries]

                            # Check if new_data contains non-integer values
                            if any(not value.isdigit() for value in new_data):
                                status_label = ttk.Label(edit_window, text='Please enter integers only', foreground='red')
                                status_label.grid(columnspan=2)
                                return

                            treeview1.insert("", "end", values=new_data)
                            edit_window.destroy()

                            data = read_csv()
                            data.append(new_data)
                            write_csv(data)

                        def edit2(edit_window, entries):

                            new_data = [entry.get() for entry in entries]
                            print(new_data, 'new data')

                            # Check if new_data contains non-integer values
                            if any(not value.isdigit() for value in new_data):
                                status_label = ttk.Label(edit_window, text='Please enter integers only', foreground='red')
                                status_label.grid(columnspan=2)
                                return

                            item_id = treeview1.focus()

                            treeview1.item(item_id, values=new_data)
                            edit_window.destroy()

                            data = read_csv()

                            print(data)
                            index = treeview1.index(item_id)
                            data[index] = new_data
                            write_csv(data)

                        def reset(entries):
                            for entry in entries:
                                entry.delete(0, 'end')

                        def on_double_click_wrapper(event):

                            selected_items = treeview1.selection()
                            if not selected_items:
                                messagebox.showerror("Error", "Please select a row!")
                                return

                            item_id = treeview1.selection()[0]
                            item = treeview1.item(item_id)

                            edit_window = tk.Toplevel(app)
                            edit_window.title("Edit Student Details")

                            topframe = tk.Frame(edit_window, padx=10, pady=10)
                            topframe.grid(row=0, column=0, sticky='nsew')

                            local_entries = []
                            for i, text in enumerate(labels_text):
                                label = ttk.Label(topframe, text=text)
                                label.grid(column=0, row=i, padx=5,pady=5, sticky=tk.W)
                                entry = ttk.Entry(topframe)
                                entry.grid(column=1, row=i, padx=5, pady=5)
                                entry.insert(0, item['values'][i])
                                local_entries.append(entry)

                            button_frame = tk.Frame(topframe, padx=5, pady=5)
                            button_frame.grid(columnspan=2, sticky='ew')

                            button_commands = [
                                lambda: save(edit_window, local_entries),   lambda: edit2(edit_window, local_entries),   lambda: reset(local_entries)
                            ]

                            for i, (text, cmd) in enumerate(zip(button_texts, button_commands)):
                                button = ttk.Button(button_frame, text=text, command=cmd)
                                button.grid(column=i, row=0, padx=5, pady=5)

                            topframe.columnconfigure(1, weight=1)

                        # Search function

                        def search():
                            search_term = search_entry.get().lower()
                            for item in treeview.get_children():
                                treeview.delete(item)

                            filtered_data = [
                                row for row in data if search_term in ' '.join(row).lower()]
                            for row in filtered_data:
                                treeview.insert("", "end", values=row)

                            total_students_label.config(text=f"Total students: {len(filtered_data)}")

                        def get_index():

                            selection = treeview1.focus()
                            index = treeview1.index(selection)
                            print(index)
                            return index

                        def delete_row():

                            selected_item = tree.focus()
                            node_name = tree.item(selected_item, 'text')
                            node_id = selected_item
                            csv_file = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                            index = get_index()

                            lines = []
                            with open(csv_file, 'r') as file:
                                reader = csv.reader(file)
                                for i, row in enumerate(reader):
                                    if i != index:
                                        lines.append(row)

                            with open(csv_file, 'w', newline='') as file:
                                writer = csv.writer(file)
                                writer.writerows(lines)

                            selected_item = treeview1.focus()
                            if selected_item:
                                treeview1.delete(selected_item)

                        # starting

                        app = tk.Toplevel()
                        app.title('Students List')

                        students_list_title = tk.Label(app, text="Students List")
                        students_list_title.grid(row=0, column=0, sticky='nsew')

                        labels_text = ["Wrong-Number", "TP ", "DS", "Exam"]
                        button_texts = ["Save", "Edit", "Reset"]
                        entries = []

                        treeview1 = ttk.Treeview(app, columns=labels_text, show="headings", selectmode="browse")

                        treeview1.grid(row=1, column=0, sticky='nsew')

                        for text in labels_text:
                            treeview1.heading(text, text=text)

                        treeview1.bind("<Double-1>", on_double_click_wrapper)

                        data = read_csv()

                        for row in data:
                            treeview1.insert("", "end", values=row)

                        total_students_label = tk.Label(app, text=f"Total students: {len(data)}")
                        total_students_label.grid(row=2, column=0, sticky='nsew')

                        app.columnconfigure(0, weight=1)
                        app.rowconfigure(0, weight=1)
                        app.rowconfigure(1, weight=1)
                        app.rowconfigure(2, weight=1)

                        # search bar bloc

                        # Create a frame for search bar
                        search_frame = tk.Frame(app)
                        search_frame.grid(row=3, column=0, sticky='nsew')

                        search_label = ttk.Label(search_frame, text="Search:")
                        search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

                        search_entry = ttk.Entry(search_frame)
                        search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')

                        search_button = ttk.Button(search_frame, text="Search", command=search)
                        search_button.grid(column=2, row=0, padx=5, pady=5)

                        search_frame.columnconfigure(1, weight=1)

                        # to define the add and delete buttons (start)

                        column_button_frame = tk.Frame(app)
                        column_button_frame.grid(row=4, column=0, sticky='ew')

                        # to define the add and delete buttons (end)

                        add_column_button = tk.Button(column_button_frame, text='Add Column', command=lambda: add_column(treeview1))
                        add_column_button.grid(column=0, row=0, padx=5, pady=5)

                        delete_column_button = tk.Button(column_button_frame, text='Delete Column', command=lambda: delete_column(treeview1))
                        delete_column_button.grid(column=1, row=0, padx=5, pady=5)

                        # index_button = tk.Button(column_button_frame, text='show index', command=lambda: get_index())
                        # index_button.grid(column=5, row=0, padx=5, pady=5)

                        openDBeditor_button = tk.Button(column_button_frame, text='open DB editor', command=lambda: on_double_click_wrapper(event=None))
                        openDBeditor_button.grid(column=2, row=0, padx=5, pady=5)

                        delete_row_button = tk.Button(column_button_frame, text='delete_row', command=lambda: delete_row())
                        delete_row_button.grid(column=3, row=0, padx=5, pady=5)

                        # activ_tree_button = tk.Button(column_button_frame, text='activ_tree', command=lambda: get_active_tree_name(trees, trees_names))
                        # activ_tree_button.grid(column=4, row=0, padx=5, pady=5)

                        # id_test_button = tk.Button(column_button_frame, text='id_test', command=lambda: get_selected_id(tree))
                        # id_test_button.grid(column=5, row=0, padx=5, pady=5)

                        # configure the window size and resize when max and min it

                        app.columnconfigure(0, weight=1)
                        app.rowconfigure(0, weight=1)
                        app.rowconfigure(1, weight=1)
                        app.rowconfigure(2, weight=1)
                        app.rowconfigure(3, weight=1)
                        app.rowconfigure(4, weight=1)

                        app.mainloop()

                    def add_column(treeview1):
                        column_name = simpledialog.askstring("Input", "Enter column name:")
                        if column_name:
                            # Fetch current columns
                            columns = list(treeview1["columns"])

                            # Append new column
                            columns.append(column_name)

                            # Re-assign the updated columns
                            treeview1['columns'] = columns

                            # Set heading and column width for the new column
                            treeview1.heading(column_name, text=column_name)
                            treeview1.column(column_name, width=100)

                            # Retain the heading and column width for the existing columns
                            for col in columns[:-1]:
                                treeview1.heading(col, text=col)
                                treeview1.column(col, width=100)

                    def delete_column(treeview1):
                        column_name = simpledialog.askstring("Input", "Enter column name to delete:")

                        if column_name:
                            column_name = column_name.strip()  # Remove leading/trailing spaces
                            if column_name == '#0':
                                messagebox.showerror("Error", "Cannot delete the 'tree' part of the Treeview.")
                                return

                            # Fetch current columns
                            columns = list(treeview1["columns"])

                            # Check if column_name exists in columns
                            if column_name not in columns:
                                messagebox.showerror("Error", f"Column '{column_name}' not found")
                                return

                            # Save current column configurations
                            column_configurations = {col: {'width': treeview1.column(col)['width'], 'text': treeview1.heading(col)[
                                                                                     'text']} for col in columns if col != column_name}

                            # Remove the column
                            columns.remove(column_name)

                            # Re-assign the updated columns
                            treeview1['columns'] = columns

                            # Apply saved configurations to existing columns
                            for col, conf in column_configurations.items():
                                treeview1.column(col, width=conf['width'])
                                treeview1.heading(col, text=conf['text'])

                            # If the deleted column is a default column, remove it from the list of default columns
                            if column_name in labels_text:
                                labels_text.remove(column_name)
                        else:
                            messagebox.showerror("Error", "No column name entered")

                    def update_to_treeview_type(tree, iid, node_id):

                        node_name = find_and_focus(tree)
                        file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}.csv"

                        # First, read the entire CSV file into a list of dictionaries
                        data = []
                        with open(file_path, mode='r', newline='', encoding='utf-8') as csvfile:
                            reader = csv.DictReader(csvfile)
                            for row in reader:
                                data.append(row)

                        # Next, update the 'Type' of the specified node
                        for row in data:
                            if row['NodeID'] == iid:
                                row['Type'] = 'treeview'

                        # Finally, write the updated data back out to the CSV file
                        with open(file_path, mode='w', newline='', encoding='utf-8') as csvfile:
                            writer = csv.DictWriter(csvfile, fieldnames=data[0].keys())
                            writer.writeheader()
                            writer.writerows(data)

                    def update_to_txt_type(tree, iid):
                        # selected_item = tree.selection()[0]
                        # node_name = tree.item(selected_item, 'text')

                        node_name = find_and_focus(tree)
                        file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}.csv"

                        # First, read the entire CSV file into a list of dictionaries
                        data = []
                        with open(file_path, mode='r', newline='', encoding='utf-8') as csvfile:
                            reader = csv.DictReader(csvfile)
                            for row in reader:
                                data.append(row)

                        # Next, update the 'Type' of the specified node

                        for row in data:
                            if row['NodeID'] == iid:
                                row['Type'] = 'txt'
                                # base = os.path.splitext(row['path'])[0]
                                # row['path'] = base + ".txt"

                        # Finally, write the updated data back out to the CSV file
                        with open(file_path, mode='w', newline='', encoding='utf-8') as csvfile:
                            writer = csv.DictWriter(csvfile, fieldnames=data[0].keys())
                            writer.writeheader()
                            writer.writerows(data)

                    def show_details(tree):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.selection()[0]
                        node_id = selected_item

                        node_name = tree.item(selected_item, 'text')

                        node_dict = read_csv_file()

                        node_type = None
                        for node in node_dict:
                            if node['NodeID'] == node_id:
                                node_type = node['Type']
                                break

                        if node_type is None:
                            print(f"No node found with ID: {node_id}")
                            return

                        print(f"Node type: {node_type}")

                        if node_type == 'treeview':
                            myDBdash(event=None, treeview_window=tree)
                        elif node_type == 'txt':
                            myED(node_id, node_name)
                        else:
                            print(f"Unknown node type: {node_type}")

                    def log_action(action, node_name):
                        with open('user_actions1.csv', 'a', newline='') as csvfile:
                            fieldnames = ['timestamp', 'action', 'node_name']
                            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                            writer.writerow({'timestamp': datetime.now(), 'action': action, 'node_name': node_name})

                    def create_txt_for_node(node_name):
                        selected_item = tree.selection()[0]
                        node_id = selected_item

                        # unique_id = str(uuid.uuid4())
                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.txt"

                        # Create a text file for the specific node
                        with open(filename, mode='w') as file:
                            file.write("This is a text file.")

                    def create_csv_for_node(node_name):
                        selected_item = tree.selection()[0]
                        node_id = selected_item
                        # unique_id = str(uuid.uuid4())
                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                        # Create the CSV file if it doesn't exist
                        if not os.path.exists(filename):
                            with open(filename, mode='w', newline='') as file:
                                writer = csv.writer(file)
                                # Write a line to the file
                                writer.writerow(["it's empty !", "(click Here)", "", "", "", "", "", "", "", ])

                    def delete_txt_for_node(node_name):
                        selected_item = tree.selection()[0]
                        node_id = selected_item

                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.txt"

                        # Delete the text file if it exists
                        if os.path.exists(filename):
                            os.remove(filename)

                    def delete_csv_for_node(node_name):

                        selected_item = tree.focus()
                        node_id = selected_item

                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                        # Delete the CSV file if it exists
                        if os.path.exists(filename):
                            os.remove(filename)

                    def myED(node_id, node_name):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        text_editor_window = tk.Toplevel()
                        text_editor_window.title("Text Editor")

                        # Create a scrolled text widget
                        text_editor = scrolledtext.ScrolledText(text_editor_window, wrap=tk.WORD, width=30, height=10)
                        text_editor.pack()

                        node_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.txt"

                        with open(node_path, 'r') as file:
                            txt_content = file.read()
                            text_editor.insert("end", txt_content)

                        # Function to save the text content and update the node dictionary
                        def save_text_content():
                            text_content = text_editor.get("1.0", tk.END)
                            with open(node_path, 'w') as file:  # open in write mode
                                file.write(text_content)

                            text_editor_window.destroy()

                        # Save button to save the text content
                        save_button = tk.Button(text_editor_window, text="Save", command=save_text_content)
                        save_button.pack()

                        # Function to save the text content and update the node dictionary

                        def save_text_content():
                            node_dict[node_id]['text_content'] = text_editor.get("1.0", tk.END)

                            # Write the text content to the associated text file
                            file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\{node_name}-{node_id}.txt"
                            with open(file_path, 'w') as file:
                                file.write(node_dict[node_id]['text_content'])

                            text_editor_window.destroy()

                    def set_node_type(tree, node_dict, node_type):

                        print("set_node_type called")  # Add this

                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.selection()[0]
                        node_name = tree.item(selected_item, 'text')
                        node_id = selected_item

                        if node_type == 'treeview':

                            delete_txt_for_node(tree.item(selected_item, 'text'))

                            create_csv_for_node(tree.item(selected_item, 'text'))

                            update_to_treeview_type(tree, iid=tree.selection()[0], node_id=node_id)

                            # myDBdash(event=None, treeview_window=tree)

                        if node_type == 'txt':

                            delete_csv_for_node(tree.item(selected_item, 'text'))

                            create_txt_for_node(tree.item(selected_item, 'text'))

                            update_to_txt_type(tree, iid=tree.selection()[0])

                            ##myED(node_id=selected_item,node_name=tree.item(selected_item, 'text'))

                    def update_dict_and_append_csv(tree, node_dict, file_path):
                        if not node_dict:
                            messagebox.showerror("Error", "No Treeview to update")
                            return

                        tree_structure = {}
                        for node_id in node_dict:
                            node_info = tree.item(node_id)
                            node_info.update(node_dict[node_id])
                            parent_id = tree.parent(node_id)
                            if parent_id:
                                node_info['parent'] = parent_id
                            tree_structure[node_id] = node_info

                        file_exists = os.path.isfile(file_path)

                        with open(file_path, 'a', newline='') as file:
                            writer = csv.writer(file)
                            if not file_exists:
                                writer.writerow(['NodeID', 'ParentID', 'Text', 'Type', 'Path'])
                            for node_id, node_info in tree_structure.items():
                                parent_id = node_info.get('parent', '')
                                node_text = node_info.get('text', '')
                                node_type = node_info.get('type', 'treeview')
                                node_path = node_info.get('path', '')
                                writer.writerow([node_id, parent_id, node_text, node_type, node_path])

                        return tree_structure

                    def save_treeview(tree, node_dict, file_path=None, override_existing=True):
                        if not node_dict:
                            messagebox.showerror("Error", "No Treeview to save")
                            return

                        tree_structure = {}
                        for node_id in node_dict:
                            node_info = tree.item(node_id)
                            node_info.update(node_dict[node_id])
                            parent_id = tree.parent(node_id)
                            if parent_id:
                                node_info['parent'] = parent_id
                            tree_structure[node_id] = node_info

                        if not file_path:
                            node_name = find_and_focus(tree)
                            file_path = os.path.join(fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2", node_name + ".csv",)

                        if override_existing and os.path.exists(file_path):
                            confirm_override = messagebox.askyesno("Override Confirmation", f"The file '{file_path}' already exists. Do you want to override it?")
                            if not confirm_override:
                                return  # User canceled the override

                        with open(file_path, 'w', newline='') as file:
                            writer = csv.writer(file)
                            writer.writerow(['NodeID', 'ParentID', 'Text', 'Type', 'Path'])
                            for node_id, node_info in tree_structure.items():
                                parent_id = node_info.get('parent', '')
                                node_text = node_info.get('text', '')
                                node_type = node_info.get('type', 'treeview')
                                node_path = node_info.get('path', '')
                                writer.writerow([node_id, parent_id, node_text, node_type, node_path])

                        log_action(f'Saved Treeview state to {file_path}', 'Treeview')
                        return file_path

                    # Note: the behavior of the code tells ttk.Treeview module to use nodes icons and nodes build-in functions
                    # and each row in the dict csv file will transform to a node

                    # Generating Nodes Process(how it works):
                    # There's a dictionary called tree_structure that contains information about nodes in a tree.
                    # Each node has an id, some text to display, parent node (if any), type of the node, and a path.
                    # A loop iterates through each node in the tree_structure dictionary.
                    # For each node, it extracts the id, parent, and text.
                    # It also creates a node_dict to store metadata about each node. (this happens in def save_treeview() FN )
                    # The tree.insert() method is used to add nodes to the tree view:
                    # The parent parameter determines where to add the node.
                    # The iid parameter sets the node's identifier.
                    # The text parameter sets the text to display for the node.
                    # This code effectively populates a tree view with nodes based on the provided dictionary, creating a hierarchical structure.

                    def load_treeview(tree, node_dict, username_entry):
                        initial_dir = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2"
                        filename = filedialog.askopenfilename(title="Select Treeview File", initialdir=initial_dir, filetypes=(("CSV Files", "*.csv"), ("All Files", "*.*")))
                        if not filename:
                            return

                        try:
                            with open(filename, 'r') as file:
                                reader = csv.reader(file)
                                next(reader)  # Skip the header row
                                tree_structure = {row[0]: {
                                    'parent': row[1], 'text': row[2], 'type': row[3], 'path': row[4]} for row in reader}
                        except FileNotFoundError:
                            messagebox.showerror("Error", "No saved Treeview state found")
                            return

                        # clearing the previous node-treeview (start)
                        for item in tree.get_children():
                            tree.delete(item)
                        node_dict.clear()

                        # clearing the previous node-treeview (end)

                        for node_id, node_info in tree_structure.items():
                            parent = node_info.get('parent', '')
                            node_text = node_info.get('text', '')
                            node_dict[node_id] = {'type': node_info.get('type', 'treeview'), 'path': node_info.get('path', '')}
                            tree.insert(parent, 'end', iid=node_id,           text=node_text)

                        csv_file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\log_data.csv"
                        data2 = []
                        # Open the CSV file in read mode
                        with open(csv_file_path, mode='r') as csvfile:
                            # Create a CSV reader object
                            reader = csv.reader(csvfile)

                            # Iterate through each row in the CSV file
                            for row in reader:
                                data2.append(row)

                        print(data2)
                        # allowed_path =''
                        for row in data2:
                            if row[5] == "user" and row[0] == username_entry.get():
                                allowed_path = row[6]
                                print(allowed_path)

                        for item in tree.get_children():
                            if tree.item(item, "text") != allowed_path:
                                tree.detach(item)

                        log_action(f'Loaded Treeview state from {filename}', 'Treeview')

                    # warning! here tree==node-tree and also tree==the treeview of logs or Login Manager

                    def add_node(tree, node_dict, parent=None):
                        node_name = simpledialog.askstring("Input", "Enter node name:")

                        if node_name is not None:
                            # Check if node name already exists
                            if node_name in [tree.item(item)['text'] for item in tree.get_children()]:
                                messagebox.showerror("Error", f"Node '{node_name}' already exists.")
                                return

                            parent = parent or ''  # If parent is None, set it to ''
                            node_id = tree.insert(parent, 'end', text=node_name)
                            node_dict[node_id] = {'type': 'treeview',       'path': f"{node_name}_data.csv"}

                            log_action('Node added', node_name)

                            print("node added!")

                            messagebox.showinfo("info", "node added!")

                    def add_subnode(tree, node_dict):
                        if not tree.selection():
                            messagebox.showerror("Error", "No parent node selected")
                            return

                        parent_node = tree.selection()[0]
                        node_name = simpledialog.askstring("Input", "Enter subnode name:")

                        if node_name is not None:
                            # Check if subnode name already exists
                            if node_name in [tree.item(item)['text'] for item in tree.get_children(parent_node)]:
                                messagebox.showerror("Error", f"Subnode '{node_name}' already exists.")
                                return

                            node_id = tree.insert(parent_node, 'end', text=node_name)
                            node_dict[node_id] = {'type': 'treeview',       'path': f"{node_name}_data.csv"}

                            log_action('Subnode added', node_name)

                            print("sub-node added!")

                            messagebox.showinfo("info", "sub-node added!")

                    def edit_node(tree, node_dict):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.selection()[0]
                        current_node_name = tree.item(selected_item)['text']

                        new_node_name = simpledialog.askstring("Input", "Enter new node name:", initialvalue=current_node_name)

                        if new_node_name is not None:
                            # Check if new node name already exists
                            if new_node_name in [tree.item(item)['text'] for item in tree.get_children() if item != selected_item]:
                                messagebox.showerror("Error", f"Node '{new_node_name}' already exists.")
                                return

                            tree.item(selected_item, text=new_node_name)
                            node_dict[selected_item]['path'] = f"{new_node_name}_data.csv"
                            log_action('Node edited', new_node_name)

                            print("node edited!")

                            messagebox.showinfo("info", "node edited!")

                    def remove_node(tree, node_dict):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.focus()
                        # Get the node name
                        node_name = tree.item(selected_item, 'text')
                        tree.delete(selected_item)
                        if selected_item in node_dict:
                            del node_dict[selected_item]
                        log_action('Node removed', selected_item)

                        delete_csv_for_node(node_name)

                        print("node removed!")

                        messagebox.showinfo("info", "node removed!")

                    def search_node(tree, search_term):
                        found = False
                        for item in tree.get_children():
                            found = search_subtree(tree, item, search_term)
                            if found:
                                break
                        if not found:
                            messagebox.showinfo("Search", "No matching node found")

                        log_action('Node searched', node_name)

                    def search_subtree(tree, parent_item, search_term):
                        found = False
                        node_text = tree.item(parent_item, 'text')
                        if re.search(search_term, node_text, re.IGNORECASE):
                            tree.selection_set(parent_item)
                            tree.focus(parent_item)
                            found = True
                        for child_item in tree.get_children(parent_item):
                            found = search_subtree(tree, child_item, search_term)
                            if found:
                                break
                        return found

                    def clear_treeview(treeview):
                        for row in treeview.get_children():
                            treeview.delete(row)

                        log_action('treeview cleared', treeview)

                    # testing function

                    def get_selected_id(tree):
                        selected_item = tree.selection()  # This method returns a tuple of selected item(s)

                        # Get the 'id' of the first selected item
                        item_id = selected_item[0]
                        print('row_id:', item_id)

                        node_name = tree.item(selected_item, 'text')
                        print('row_name:', node_name)

                    # testing function

                    def find_and_focus(tree):

                        # Get all items
                        items = tree.get_children('')

                        # Check if there are items in the tree
                        if items:
                            # Get the first item (which is typically the root in a treeview)
                            top_node = items[0]

                            # Get the data for the top node
                            top_node_data = tree.item(top_node)

                            # Get the 'text' attribute from the top node data, which is the name of the node
                            top_node_name = top_node_data.get('text', '')

                            # return print(f"Top node name: {top_node_name}")
                            # node_name=top_node_name

                            return top_node_name

                        else:
                            print("No items in the treeview.")
                            return None

                    def read_csv_file():

                        node_name = find_and_focus(tree)

                        file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}.csv"

                        with open(file_path, 'r') as csvfile:
                            reader = csv.DictReader(csvfile)
                            node_dict = list(reader)
                        return node_dict

                    # it dispaly the name of the selected treeview (just for testing: to know which treeview we working on )
                ##
                    # def on_select(event):
                    # print("Selection:", event.widget.selection())
                    ##
                    # treeview1.bind("<<TreeviewSelect>>", on_select)

                    def check_and_write(file_path):
                        # Check if the file is empty
                        if os.path.getsize(file_path) == 0:
                            # Open the file in write mode
                            with open(file_path, 'w', newline='') as csvfile:
                                # Create a writer object
                                writer = csv.writer(csvfile)
                                # Write a line to the file
                                writer.writerow(["it's empty"])

                 
                    ### a testing function

                    def open_webpage():

                          webbrowser.open(  'http://localhost/Libraryms-PHP/librarian/home.php')

                    def on_save(event):
                                        """Saves the treeview state when the user presses `Ctrl`+`S`."""
                                        save_treeview(tree, node_dict, file_path=None, override=True)

                                        tree.bind("<Control-S>", on_save)


                    def set_allowed_path(tree):  # the logs "tree" is visible and not "tree" that is node-treeview
                        item = tree.focus()  # selection()[0]

                        folder = tree.item(item, "text")
                        print(folder)

                        win = tk.Tk()
                        win.title("allow a path to a user")
                        win.geometry("200x100")
                        user_label = tk.Label(win, text="Username:")
                        user_label.pack()
                        user_entry = tk.Entry(win)
                        user_entry.pack()
                        user_entry_button = tk.Button(win, text="set_path", command=lambda: set_path (item, user_entry, folder))
                        user_entry_button.pack()
                        return (user_entry, folder)


                    def set_path(tree, user_entry, folder):
                        user = user_entry.get()
                        print(user)

                        # Specify the path to your CSV file
                        csv_file_path = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\log_data.csv'

                        data2 = []
                        # Open the CSV file in read mode
                        with open(csv_file_path, mode='r') as csvfile:
                            # Create a CSV reader object
                            reader = csv.reader(csvfile)

                            # Iterate through each row in the CSV file
                            for row in reader:
                                data2.append(row)

                            print("data2", data2)
                        for i, row in enumerate(data2):
                            # Process each row (row is a list of values)
                            if row[0] == user and row[5] == "user":
                                # Perform actions here based on the condition
                                row[6] = folder
                                index = i
                                print('row index:', i)
                                print("row6:", row[6])
                                break

                        data2_cell = data2[index][6] = folder

                        print("data2 call", data2_cell)

                        with open(csv_file_path, mode='w', newline='') as csvfile:
                            # Create a CSV writer object
                            writer = csv.writer(csvfile)
                            for row in data2:

                                writer.writerow(row)

                        print("data2", data2)

                        print("Allowed Path has been set!")

                        messagebox.showinfo("info", "Allowed Path has been set!")

                        log_action('Allowed Path has been set!', item)
                        
                    def history():
                        file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\history_file.csv"
                        try:
                            with open(file_path, 'r') as file:
                                content = file.read()
                            content_window = tk.Toplevel(root)
                            content_window.title("All Records")
                            content_window.geometry("600x400")
                            content_window.configure(bg='#00B0F0')
                            text_widget = ScrolledText(content_window, wrap=tk.WORD, font=("Helvetica", 12), bg='white')
                            text_widget.insert(tk.END, content)
                            text_widget.pack(fill=tk.BOTH, expand=True)
                            print("Hisdtory imported!")
                            messagebox.showinfo("info", "Hisdtory imported!")
                            log_action('Hisdtory imported!', file_path)
                        except FileNotFoundError:
                            tk.messagebox.showerror("File Not Found", "The specified file does not exist.")
                        except Exception as e:
                            tk.messagebox.showerror("Error", f"An error occurred: {e}")
                    # starting

                    node_dict = {}

                    add_button = tk.Button(root, text='Add Node',command=lambda: add_node(tree, node_dict, None))
                    add_button.pack(side="left", padx=10, pady=5)

                    add_subnode_button = tk.Button(root, text='Add Subnode', command=lambda: add_subnode(tree, node_dict))
                    add_subnode_button.pack(side="left", padx=10, pady=5)

                    edit_button = tk.Button(root, text='Edit Node', command=lambda: edit_node(tree, node_dict))
                    edit_button.pack(side="left", padx=10, pady=5)

                    delete_button = tk.Button(root, text='Delete Node',command=lambda: remove_node(tree, node_dict))
                    delete_button.pack(side="left", padx=10, pady=5)

                    detail_button = tk.Button(root, text='Show Details',command=lambda: show_details(tree))  # , node_dict
                    detail_button.pack(side="left", padx=10, pady=5)

                    treeview_button = tk.Button(root, text='Set Treeview',command=lambda: set_node_type(tree, node_dict, 'treeview'))
                    treeview_button.pack(side="left", padx=10, pady=5)

                    txt_button = tk.Button(root, text='Set TXT',command=lambda: set_node_type(tree, node_dict, node_type='txt'))
                    txt_button.pack(side="left", padx=10, pady=5)

                    save_button = tk.Button(root, text='Save Node-Treeview\n cltr + S ', command=lambda: save_treeview(tree, node_dict, file_path=None, override_existing=True))
                    save_button.pack(side="left", padx=10, pady=5)

                    load_button = tk.Button(root, text='Load Treeview',command=lambda: load_treeview(tree, node_dict, username_entry))
                    load_button.pack(side="left", padx=10, pady=5)

                    search_button = tk.Button(root, text='Search', command=lambda: search_node(tree, simpledialog.askstring("Input", "Enter search term:")))
                    search_button.pack(side="left", padx=10, pady=5)

                    # activ_tree_button = tk.Button(root, text='activ_tree', command=lambda: get_active_tree_name(trees, trees_names))
                    # activ_tree_button.pack(side="left", padx=10, pady=5)

                ##
                    # id_test_button = tk.Button(root, text='id_test', command=lambda: get_selected_id(tree))
                    # id_test_button.pack(side="left", padx=10, pady=5)

                    # focus_button = tk.Button(root, text='focus', command=lambda: find_and_focus(tree))
                    # focus_button.pack(side="left", padx=10, pady=5)

                    # button = tk.Button(root, text='Open Web Page', command=open_webpage)
                    # button.pack(side="left", padx=10, pady=5)

                    history_button = tk.Button(root, text="History", command=history)
                    history_button.pack(side="left", padx=10, pady=5)
                    # history_button.place(x=10, y=10)

                    set_path_button = tk.Button(root, text="Set Allowed Path", command=lambda: set_allowed_path(tree))
                    set_path_button.pack(side="left", padx=10, pady=5)
                    # set_path_button.place(x=100, y=10)

                    # Assume 'treeview' is the Treeview widget
                    clear_button = tk.Button(root, text="Clear", command=lambda: clear_treeview(tree))
                    clear_button.pack(side="left", padx=10, pady=5)


        def user_section():

 
                    print ("u entred usesr section")

                    class NodeDetailWindow(tk.Toplevel):
                        def __init__(self, master, node_name, node_type, node_path):
                            super().__init__(master)
                            self.title("Node Detail")

                            # Create labels and display node details
                            name_label = tk.Label(self, text="Node Name:")
                            name_label.pack()
                            name_value = tk.Label(self, text=node_name)
                            name_value.pack()

                            type_label = tk.Label(self, text="Node Type:")
                            type_label.pack()
                            type_value = tk.Label(self, text=node_type)
                            type_value.pack()

                            path_label = tk.Label(self, text="Node Path:")
                            path_label.pack()
                            path_value = tk.Label(self, text=node_path)
                            path_value.pack()

                    root = tk.Tk()
                    root.title("Treeview Syst")

                    tree = ttk.Treeview(root)
                    tree.pack(side="top", fill="both", expand=True)

                    labels_text = ["Wrong-Number", "TP", "DS", "Exam"]

                    treeview1 = ttk.Treeview(root, columns=labels_text, show="headings", selectmode="browse")

                    def write_data_to_node(node_name, data):
                        filename = f"{node_name}.csv"

                        # Write data to the specific node's CSV file
                        with open(filename, mode='a', newline='') as file:
                            writer = csv.writer(file)
                            writer.writerow(data)

                    def get_active_tree_name(trees, trees_names):
                            active_tree = None
                            for tree in trees:
                                if tree.focus():
                                    active_tree = tree
                                    break

                            if active_tree is None:
                                print("No active tree found.")
                                return None

                            return print(trees_names[active_tree] )

                    # tree (is node-treeview ) (this is the only active treeview ,   "tree=t2")

                    # treeview1 == simple treeview inside

                    def myDBdash(event, treeview_window):

                        def validate_input(new_value):
                            if new_value == "":
                                return True
                            try:
                                int(new_value)
                                error_label.grid_forget()  # Hide the error message if input is valid
                                return True
                            except ValueError:
                                # Show the error message below the entry
                                error_label.grid(column=1, row=i+1, padx=5, pady=2)
                                return False

                        selected_item = treeview_window.focus()
                        if selected_item:
                            node_name = tree.item(selected_item, 'text')

                        selected_item = tree.selection()[0]
                        node_id = selected_item
                        csv_file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                        def read_csv():

                            selected_item = tree.focus()
                            node_name = tree.item(selected_item, 'text')
                            node_id = selected_item
                            csv_file = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                            data = []
                            with open(csv_file, mode='r') as csvfile:
                                reader = csv.reader(csvfile)
                                for row in reader:
                                    data.append(row)
                            return data

                        def write_csv(data):

                            selected_item = tree.focus()
                            node_name = tree.item(selected_item, 'text')
                            node_id = selected_item
                            csv_file = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                            with open(csv_file, mode='w', newline='', encoding='utf-8') as csvfile:
                                writer = csv.writer(csvfile)
                                for row in data:
                                    writer.writerow(row)

                        def save(edit_window, entries):
                            new_data = [entry.get() for entry in entries]

                            # Check if new_data contains non-integer values
                            if any(not value.isdigit() for value in new_data):
                                status_label = ttk.Label(edit_window, text='Please enter integers only', foreground='red')
                                status_label.grid(columnspan=2)
                                return

                            treeview1.insert("", "end", values=new_data)
                            edit_window.destroy()

                            data = read_csv()
                            data.append(new_data)
                            write_csv(data)

                        def edit2(edit_window, entries):

                            new_data = [entry.get() for entry in entries]
                            print(new_data, 'new data')

                            # Check if new_data contains non-integer values
                            if any(not value.isdigit() for value in new_data):
                                status_label = ttk.Label(edit_window, text='Please enter integers only', foreground='red')
                                status_label.grid(columnspan=2)
                                return

                            item_id = treeview1.focus()

                            treeview1.item(item_id, values=new_data)
                            edit_window.destroy()

                            data = read_csv()

                            print(data)
                            index = treeview1.index(item_id)
                            data[index] = new_data
                            write_csv(data)

                        def reset(entries):
                            for entry in entries:
                                entry.delete(0, 'end')

                        def on_double_click_wrapper(event):

                            selected_items = treeview1.selection()
                            if not selected_items:
                                messagebox.showerror("Error", "Please select a row!")
                                return

                            item_id = treeview1.selection()[0]
                            item = treeview1.item(item_id)

                            edit_window = tk.Toplevel(app)
                            edit_window.title("Edit Student Details")

                            topframe = tk.Frame(edit_window, padx=10, pady=10)
                            topframe.grid(row=0, column=0, sticky='nsew')

                            local_entries = []
                            for i, text in enumerate(labels_text):
                                label = ttk.Label(topframe, text=text)
                                label.grid(column=0, row=i, padx=5,pady=5, sticky=tk.W)
                                entry = ttk.Entry(topframe)
                                entry.grid(column=1, row=i, padx=5, pady=5)
                                entry.insert(0, item['values'][i])
                                local_entries.append(entry)

                            button_frame = tk.Frame(topframe, padx=5, pady=5)
                            button_frame.grid(columnspan=2, sticky='ew')

                            button_commands = [
                                lambda: save(edit_window, local_entries),   lambda: edit2(edit_window, local_entries),   lambda: reset(local_entries)
                            ]

                            for i, (text, cmd) in enumerate(zip(button_texts, button_commands)):
                                button = ttk.Button(button_frame, text=text, command=cmd)
                                button.grid(column=i, row=0, padx=5, pady=5)

                            topframe.columnconfigure(1, weight=1)

                        # Search function


                        def search():
                            search_term = search_entry.get().lower()
                            for item in treeview.get_children():
                                treeview.delete(item)

                            filtered_data = [
                                row for row in data if search_term in ' '.join(row).lower()]
                            for row in filtered_data:
                                treeview.insert("", "end", values=row)

                            total_students_label.config(text=f"Total students: {len(filtered_data)}")

                        def get_index():

                            selection = treeview1.focus()
                            index = treeview1.index(selection)
                            print(index)
                            return index

                        def delete_row():

                            selected_item = tree.focus()
                            node_name = tree.item(selected_item, 'text')
                            node_id = selected_item
                            csv_file = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                            index = get_index()

                            lines = []
                            with open(csv_file, 'r') as file:
                                reader = csv.reader(file)
                                for i, row in enumerate(reader):
                                    if i != index:
                                        lines.append(row)

                            with open(csv_file, 'w', newline='') as file:
                                writer = csv.writer(file)
                                writer.writerows(lines)


                            selected_item = treeview1.focus()
                            if selected_item:
                                treeview1.delete(selected_item)

                        # starting


                        app = tk.Toplevel()
                        app.title('Students List')

                        students_list_title = tk.Label(app, text="Students List")
                        students_list_title.grid(row=0, column=0, sticky='nsew')

                        labels_text = ["Wrong-Number", "TP ", "DS", "Exam"]
                        button_texts = ["Save", "Edit", "Reset"]
                        entries = []


                        treeview1 = ttk.Treeview(app, columns=labels_text, show="headings", selectmode="browse")

                        treeview1.grid(row=1, column=0, sticky='nsew')

                        for text in labels_text:
                            treeview1.heading(text, text=text)

                        treeview1.bind("<Double-1>", on_double_click_wrapper)

                        data = read_csv()

                        for row in data:
                            treeview1.insert("", "end", values=row)

                        total_students_label = tk.Label(app, text=f"Total students: {len(data)}")
                        total_students_label.grid(row=2, column=0, sticky='nsew')

                        app.columnconfigure(0, weight=1)
                        app.rowconfigure(0, weight=1)
                        app.rowconfigure(1, weight=1)
                        app.rowconfigure(2, weight=1)

                        # search bar bloc

                        # Create a frame for search bar
                        search_frame = tk.Frame(app)
                        search_frame.grid(row=3, column=0, sticky='nsew')

                        search_label = ttk.Label(search_frame, text="Search:")
                        search_label.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

                        search_entry = ttk.Entry(search_frame)
                        search_entry.grid(column=1, row=0, padx=5, pady=5, sticky='ew')

                        search_button = ttk.Button(search_frame, text="Search", command=search)
                        search_button.grid(column=2, row=0, padx=5, pady=5)

                        search_frame.columnconfigure(1, weight=1)

                        # to define the add and delete buttons (start)

                        column_button_frame = tk.Frame(app)
                        column_button_frame.grid(row=4, column=0, sticky='ew')

                        # to define the add and delete buttons (end)

                        add_column_button = tk.Button(column_button_frame, text='Add Column', command=lambda: add_column(treeview1))
                        add_column_button.grid(column=0, row=0, padx=5, pady=5)

                        delete_column_button = tk.Button(column_button_frame, text='Delete Column', command=lambda: delete_column(treeview1))
                        delete_column_button.grid(column=1, row=0, padx=5, pady=5)

                        # index_button = tk.Button(column_button_frame, text='show index', command=lambda: get_index())
                        # index_button.grid(column=5, row=0, padx=5, pady=5)

                        openDBeditor_button = tk.Button(column_button_frame, text='open DB editor', command=lambda: on_double_click_wrapper(event=None))
                        openDBeditor_button.grid(column=2, row=0, padx=5, pady=5)

                        delete_row_button = tk.Button(column_button_frame, text='delete_row', command=lambda: delete_row())
                        delete_row_button.grid(column=3, row=0, padx=5, pady=5)

                        # activ_tree_button = tk.Button(column_button_frame, text='activ_tree', command=lambda: get_active_tree_name(trees, trees_names))
                        # activ_tree_button.grid(column=4, row=0, padx=5, pady=5)

                        # id_test_button = tk.Button(column_button_frame, text='id_test', command=lambda: get_selected_id(tree))
                        # id_test_button.grid(column=5, row=0, padx=5, pady=5)


                        # configure the window size and resize when max and min it

                        app.columnconfigure(0, weight=1)
                        app.rowconfigure(0, weight=1)
                        app.rowconfigure(1, weight=1)
                        app.rowconfigure(2, weight=1)
                        app.rowconfigure(3, weight=1)
                        app.rowconfigure(4, weight=1)

                        app.mainloop()

                    def add_column(treeview1):
                        column_name = simpledialog.askstring("Input", "Enter column name:")
                        if column_name:
                            # Fetch current columns
                            columns = list(treeview1["columns"])

                            # Append new column
                            columns.append(column_name)

                            # Re-assign the updated columns
                            treeview1['columns'] = columns

                            # Set heading and column width for the new column
                            treeview1.heading(column_name, text=column_name)
                            treeview1.column(column_name, width=100)

                            # Retain the heading and column width for the existing columns
                            for col in columns[:-1]:
                                treeview1.heading(col, text=col)
                                treeview1.column(col, width=100)

                    def delete_column(treeview1):
                        column_name = simpledialog.askstring("Input", "Enter column name to delete:")

                        if column_name:
                            column_name = column_name.strip()  # Remove leading/trailing spaces
                            if column_name == '#0':
                                messagebox.showerror("Error", "Cannot delete the 'tree' part of the Treeview.")
                                return

                            # Fetch current columns
                            columns = list(treeview1["columns"])

                            # Check if column_name exists in columns
                            if column_name not in columns:
                                messagebox.showerror("Error", f"Column '{column_name}' not found")
                                return

                            # Save current column configurations
                            column_configurations = {col: {'width': treeview1.column(col)['width'], 'text': treeview1.heading(col)[
                                                                                     'text']} for col in columns if col != column_name}

                            # Remove the column
                            columns.remove(column_name)

                            # Re-assign the updated columns
                            treeview1['columns'] = columns

                            # Apply saved configurations to existing columns
                            for col, conf in column_configurations.items():
                                treeview1.column(col, width=conf['width'])
                                treeview1.heading(col, text=conf['text'])

                            # If the deleted column is a default column, remove it from the list of default columns
                            if column_name in labels_text:
                                labels_text.remove(column_name)
                        else:
                            messagebox.showerror("Error", "No column name entered")


                    def update_to_treeview_type(tree, iid, node_id):


                        node_name = find_and_focus(tree)
                        file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}.csv"

                        # First, read the entire CSV file into a list of dictionaries
                        data = []
                        with open(file_path, mode='r', newline='', encoding='utf-8') as csvfile:
                            reader = csv.DictReader(csvfile)
                            for row in reader:
                                data.append(row)

                        # Next, update the 'Type' of the specified node
                        for row in data:
                            if row['NodeID'] == iid:
                                row['Type'] = 'treeview'

                        # Finally, write the updated data back out to the CSV file
                        with open(file_path, mode='w', newline='', encoding='utf-8') as csvfile:
                            writer = csv.DictWriter(csvfile, fieldnames=data[0].keys())
                            writer.writeheader()
                            writer.writerows(data)

                    def update_to_txt_type(tree, iid):
                        # selected_item = tree.selection()[0]
                        # node_name = tree.item(selected_item, 'text')

                        node_name = find_and_focus(tree)
                        file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}.csv"

                        # First, read the entire CSV file into a list of dictionaries
                        data = []
                        with open(file_path, mode='r', newline='', encoding='utf-8') as csvfile:
                            reader = csv.DictReader(csvfile)
                            for row in reader:
                                data.append(row)

                        # Next, update the 'Type' of the specified node

                        for row in data:
                            if row['NodeID'] == iid:
                                row['Type'] = 'txt'
                                # base = os.path.splitext(row['path'])[0]
                                # row['path'] = base + ".txt"

                        # Finally, write the updated data back out to the CSV file
                        with open(file_path, mode='w', newline='', encoding='utf-8') as csvfile:
                            writer = csv.DictWriter(csvfile, fieldnames=data[0].keys())
                            writer.writeheader()
                            writer.writerows(data)

                    def show_details(tree):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.selection()[0]
                        node_id = selected_item

                        node_name = tree.item(selected_item, 'text')

                        node_dict = read_csv_file()

                        node_type = None
                        for node in node_dict:
                            if node['NodeID'] == node_id:
                                node_type = node['Type']
                                break

                        if node_type is None:
                            print(f"No node found with ID: {node_id}")
                            return

                        print(f"Node type: {node_type}")

                        if node_type == 'treeview':
                            myDBdash(event=None, treeview_window=tree)
                        elif node_type == 'txt':
                            myED(node_id, node_name)
                        else:
                            print(f"Unknown node type: {node_type}")

                    def log_action(action, node_name):
                        with open('user_actions1.csv', 'a', newline='') as csvfile:
                            fieldnames = ['timestamp', 'action', 'node_name']
                            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                            writer.writerow({'timestamp': datetime.now(), 'action': action, 'node_name': node_name})

                    def create_txt_for_node(node_name):
                        selected_item = tree.selection()[0]
                        node_id = selected_item

                        # unique_id = str(uuid.uuid4())
                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.txt"

                        # Create a text file for the specific node
                        with open(filename, mode='w') as file:
                            file.write("This is a text file.")

                    def create_csv_for_node(node_name):
                        selected_item = tree.selection()[0]
                        node_id = selected_item
                        # unique_id = str(uuid.uuid4())
                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                        # Create the CSV file if it doesn't exist
                        if not os.path.exists(filename):
                            with open(filename, mode='w', newline='') as file:
                                writer = csv.writer(file)
                                # Write a line to the file
                                writer.writerow(["it's empty !", "(click Here)", "","","","","","","",])

                    def delete_txt_for_node(node_name):
                        selected_item = tree.selection()[0]
                        node_id = selected_item

                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.txt"

                        # Delete the text file if it exists
                        if os.path.exists(filename):
                            os.remove(filename)

                    def delete_csv_for_node(node_name):

                        selected_item = tree.focus()
                        node_id = selected_item

                        filename = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.csv"

                        # Delete the CSV file if it exists
                        if os.path.exists(filename):
                            os.remove(filename)


                    def myED(node_id, node_name ):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        text_editor_window = tk.Toplevel()
                        text_editor_window.title("Text Editor")

                        # Create a scrolled text widget
                        text_editor = scrolledtext.ScrolledText(text_editor_window, wrap=tk.WORD, width=30, height=10)
                        text_editor.pack()


                        node_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}-{node_id}.txt"

                        with open(node_path, 'r') as file:
                            txt_content = file.read()
                            text_editor.insert("end", txt_content)

                        # Function to save the text content and update the node dictionary
                        def save_text_content():
                            text_content = text_editor.get("1.0", tk.END)
                            with open(node_path, 'w') as file:  # open in write mode
                                file.write(text_content)

                            text_editor_window.destroy()

                        # Save button to save the text content
                        save_button = tk.Button(text_editor_window, text="Save", command=save_text_content)
                        save_button.pack()

                        # Function to save the text content and update the node dictionary

                        def save_text_content():
                            node_dict[node_id]['text_content'] = text_editor.get("1.0", tk.END)

                            # Write the text content to the associated text file
                            file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\{node_name}-{node_id}.txt"
                            with open(file_path, 'w') as file:
                                file.write(node_dict[node_id]['text_content'])

                            text_editor_window.destroy()

                    def set_node_type(tree, node_dict, node_type):

                        print("set_node_type called")  # Add this

                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.selection()[0]
                        node_name = tree.item(selected_item, 'text')
                        node_id = selected_item


                        if node_type == 'treeview':

                            delete_txt_for_node(tree.item(selected_item, 'text'))

                            create_csv_for_node(tree.item(selected_item, 'text'))

                            update_to_treeview_type(tree, iid=tree.selection()[0], node_id=node_id)

                            # myDBdash(event=None, treeview_window=tree)


                        if node_type == 'txt':

                            delete_csv_for_node(tree.item(selected_item, 'text'))

                            create_txt_for_node(tree.item(selected_item, 'text'))

                            update_to_txt_type(tree, iid=tree.selection()[0])

                            ##myED(node_id= selected_item, node_name=tree.item(selected_item, 'text'))

                    def update_dict_and_append_csv(tree, node_dict, file_path):
                        if not node_dict:
                            messagebox.showerror("Error", "No Treeview to update")
                            return

                        tree_structure = {}
                        for node_id in node_dict:
                            node_info = tree.item(node_id)
                            node_info.update(node_dict[node_id])
                            parent_id = tree.parent(node_id)
                            if parent_id:
                                node_info['parent'] = parent_id
                            tree_structure[node_id] = node_info

                        file_exists = os.path.isfile(file_path)

                        with open(file_path, 'a', newline='') as file:
                            writer = csv.writer(file)
                            if not file_exists:
                                writer.writerow(['NodeID', 'ParentID', 'Text', 'Type', 'Path'])
                            for node_id, node_info in tree_structure.items():
                                parent_id = node_info.get('parent', '')
                                node_text = node_info.get('text', '')
                                node_type = node_info.get('type', 'treeview')
                                node_path = node_info.get('path', '')
                                writer.writerow([node_id, parent_id, node_text, node_type, node_path])

                        return tree_structure

                    def save_treeview(tree, node_dict, file_path=None, override_existing=True):
                        if not node_dict:
                            messagebox.showerror("Error", "No Treeview to save")
                            return

                        tree_structure = {}
                        for node_id in node_dict:
                            node_info = tree.item(node_id)
                            node_info.update(node_dict[node_id])
                            parent_id = tree.parent(node_id)
                            if parent_id:
                                node_info['parent'] = parent_id
                            tree_structure[node_id] = node_info

                        if not file_path:
                            node_name = find_and_focus(tree)
                            file_path = os.path.join(fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2", node_name + ".csv",)

                        if override_existing and os.path.exists(file_path):
                            confirm_override = messagebox.askyesno("Override Confirmation", f"The file '{file_path}' already exists. Do you want to override it?")
                            if not confirm_override:
                                return  # User canceled the override

                        with open(file_path, 'w', newline='') as file:
                            writer = csv.writer(file)
                            writer.writerow(['NodeID', 'ParentID', 'Text', 'Type', 'Path'])
                            for node_id, node_info in tree_structure.items():
                                parent_id = node_info.get('parent', '')
                                node_text = node_info.get('text', '')
                                node_type = node_info.get('type', 'treeview')
                                node_path = node_info.get('path', '')
                                writer.writerow([node_id, parent_id, node_text, node_type, node_path])

                        log_action(f'Saved Treeview state to {file_path}', 'Treeview')
                        return file_path

                    def load_treeview(tree, node_dict, username_entry):
                        initial_dir = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2"
                        filename = filedialog.askopenfilename(title="Select Treeview File", initialdir=initial_dir, filetypes=(("CSV Files", "*.csv"), ("All Files", "*.*")))
                        if not filename:
                            return

                        try:
                            with open(filename, 'r') as file:
                                reader = csv.reader(file)
                                next(reader)  # Skip the header row
                                tree_structure = {row[0]: {
                                    'parent': row[1], 'text': row[2], 'type': row[3], 'path': row[4]} for row in reader}
                        except FileNotFoundError:
                            messagebox.showerror("Error", "No saved Treeview state found")
                            return

                        # clearing the previous node-treeview (start)
                        for item in tree.get_children():
                            tree.delete(item)
                        node_dict.clear()

                        # clearing the previous node-treeview (end)

                        for node_id, node_info in tree_structure.items():
                            parent = node_info.get('parent', '')
                            node_text = node_info.get('text', '')
                            node_dict[node_id] = {'type': node_info.get('type', 'treeview'), 'path': node_info.get('path', '')}
                            tree.insert(parent, 'end', iid=node_id,           text=node_text)


                        csv_file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\log_data.csv"
                        data2 = []
                        # Open the CSV file in read mode
                        with open(csv_file_path, mode='r') as csvfile:
                            # Create a CSV reader object
                            reader = csv.reader(csvfile)

                            # Iterate through each row in the CSV file
                            for row in reader:
                                data2.append(row)

                        print(data2)
                        # allowed_path =''
                        for row in data2:
                            if row[5] == "user" and row[0]==username_entry.get():         
                                allowed_path = row[6]
                                print(allowed_path)

                        for item in tree.get_children():
                            if tree.item(item, "text") != allowed_path:
                                tree.detach(item)

                        log_action(f'Loaded Treeview state from {filename}', 'Treeview')

                    # warning! here tree==node-tree and also tree==the treeview of logs or Login Manager

                    def add_node(tree, node_dict, parent=None):
                        node_name = simpledialog.askstring("Input", "Enter node name:")

                        if node_name is not None:
                            # Check if node name already exists
                            if node_name in [tree.item(item)['text'] for item in tree.get_children()]:
                                messagebox.showerror("Error", f"Node '{node_name}' already exists.")
                                return

                            parent = parent or ''  # If parent is None, set it to ''
                            node_id = tree.insert(parent, 'end', text=node_name)
                            node_dict[node_id] = {'type': 'treeview',       'path': f"{node_name}_data.csv"}

                            log_action('Node added', node_name)

                            print("node added!")

                            messagebox.showinfo("info", "node added!")

                    def add_subnode(tree, node_dict):
                        if not tree.selection():
                            messagebox.showerror("Error", "No parent node selected")
                            return

                        parent_node = tree.selection()[0]
                        node_name = simpledialog.askstring("Input", "Enter subnode name:")

                        if node_name is not None:
                            # Check if subnode name already exists
                            if node_name in [tree.item(item)['text'] for item in tree.get_children(parent_node)]:
                                messagebox.showerror("Error", f"Subnode '{node_name}' already exists.")
                                return

                            node_id = tree.insert(parent_node, 'end', text=node_name)
                            node_dict[node_id] = {'type': 'treeview',       'path': f"{node_name}_data.csv"}

                            log_action('Subnode added', node_name)

                            print("sub-node added!")

                            messagebox.showinfo("info", "sub-node added!")

                    def edit_node(tree, node_dict):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.selection()[0]
                        current_node_name = tree.item(selected_item)['text']

                        new_node_name = simpledialog.askstring("Input", "Enter new node name:", initialvalue=current_node_name)

                        if new_node_name is not None:
                            # Check if new node name already exists
                            if new_node_name in [tree.item(item)['text'] for item in tree.get_children() if item != selected_item]:
                                messagebox.showerror("Error", f"Node '{new_node_name}' already exists.")
                                return

                            tree.item(selected_item, text=new_node_name)
                            node_dict[selected_item]['path'] = f"{new_node_name}_data.csv"
                            log_action('Node edited', new_node_name)

                            print("node edited!")

                            messagebox.showinfo("info", "node edited!")

                    def remove_node(tree, node_dict):
                        if not tree.selection():
                            messagebox.showerror("Error", "No node selected")
                            return

                        selected_item = tree.focus()
                        # Get the node name
                        node_name = tree.item(selected_item, 'text')
                        tree.delete(selected_item)
                        if selected_item in node_dict:
                            del node_dict[selected_item]
                        log_action('Node removed', selected_item)

                        delete_csv_for_node(node_name)

                        print("node removed!")

                        messagebox.showinfo("info", "node removed!")

                    def search_node(tree, search_term):
                        found = False
                        for item in tree.get_children():
                            found = search_subtree(tree, item, search_term)
                            if found:
                                break
                        if not found:
                            messagebox.showinfo("Search", "No matching node found")

                        log_action('Node searched', node_name)

                    def search_subtree(tree, parent_item, search_term):
                        found = False
                        node_text = tree.item(parent_item, 'text')
                        if re.search(search_term, node_text, re.IGNORECASE):
                            tree.selection_set(parent_item)
                            tree.focus(parent_item)
                            found = True
                        for child_item in tree.get_children(parent_item):
                            found = search_subtree(tree, child_item, search_term)
                            if found:
                                break
                        return found

                    def clear_treeview(treeview):
                        for row in treeview.get_children():
                            treeview.delete(row)

                        log_action('treeview cleared', treeview)

                    # testing function


                    def get_selected_id(tree):
                        selected_item = tree.selection()  # This method returns a tuple of selected item(s)

                        # Get the 'id' of the first selected item
                        item_id = selected_item[0]
                        print ('row_id:', item_id)

                        node_name = tree.item(selected_item, 'text')
                        print('row_name:', node_name)

                    # testing function


                    def find_and_focus(tree):

                        # Get all items
                        items = tree.get_children('')

                        # Check if there are items in the tree
                        if items:
                            # Get the first item (which is typically the root in a treeview)
                            top_node = items[0]

                            # Get the data for the top node
                            top_node_data = tree.item(top_node)

                            # Get the 'text' attribute from the top node data, which is the name of the node
                            top_node_name = top_node_data.get('text', '')

                            # return print(f"Top node name: {top_node_name}")
                            # node_name=top_node_name

                            return top_node_name

                        else:
                            print("No items in the treeview.")
                            return None

                    def read_csv_file():

                        node_name = find_and_focus(tree)

                        file_path = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\trees2\{node_name}.csv"

                        with open(file_path, 'r') as csvfile:
                            reader = csv.DictReader(csvfile)
                            node_dict = list(reader)
                        return node_dict

                    # it dispaly the name of the selected treeview (just for testing: to know which treeview we working on )
                ##
                    # def on_select(event):
                    # print("Selection:", event.widget.selection())
                    ##
                    # treeview1.bind("<<TreeviewSelect>>", on_select)

                    def check_and_write(file_path):
                        # Check if the file is empty
                        if os.path.getsize(file_path) == 0:
                            # Open the file in write mode
                            with open(file_path, 'w', newline='') as csvfile:
                                # Create a writer object
                                writer = csv.writer(csvfile)
                                # Write a line to the file
                                writer.writerow(["it's empty"])




                    # a testing function


                    ##def open_webpage():
##
                    ##      webbrowser.open('http://localhost/Libraryms-PHP/librarian/home.php')
##
                    # def on_save(event):
                    # """Saves the treeview state when the user presses `Ctrl`+`S`."""
                    # save_treeview(tree, node_dict, file_path=None, override=True)
##
                    # tree.bind("<Control-S>", on_save) 


                    def set_allowed_path(tree):  # the logs "tree" is visible and not "tree" that is node-treeview
                        item = tree.focus()  # selection()[0]

                        folder = tree.item(item, "text")
                        print(folder)

                        win = tk.Tk()
                        win.title("allow a path to a user")
                        win.geometry("200x100")
                        user_label = tk.Label(win, text="Username:")
                        user_label.pack()
                        user_entry = tk.Entry(win)
                        user_entry.pack()
                        user_entry_button = tk.Button(win, text="set_path", command=lambda: set_path (item, user_entry, folder))
                        user_entry_button.pack()
                        return (user_entry, folder)


                    def set_path(tree, user_entry, folder):
                        user = user_entry.get()
                        print(user)

                        # Specify the path to your CSV file
                        csv_file_path = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\log_data.csv'

                        data2 = []
                        # Open the CSV file in read mode
                        with open(csv_file_path, mode='r') as csvfile:
                            # Create a CSV reader object
                            reader = csv.reader(csvfile)

                            # Iterate through each row in the CSV file
                            for row in reader:
                                data2.append(row)

                            print("data2", data2)
                        for i, row in enumerate(data2):
                            # Process each row (row is a list of values)
                            if row[0] == user and row[5] == "user":
                                # Perform actions here based on the condition
                                row[6] = folder
                                index = i
                                print('row index:', i)
                                print("row6:", row[6])
                                break

                        data2_cell = data2[index][6] = folder

                        print("data2 call", data2_cell)

                        with open(csv_file_path, mode='w', newline='') as csvfile:
                            # Create a CSV writer object
                            writer = csv.writer(csvfile)
                            for row in data2:

                                writer.writerow(row)

                        print("data2", data2)

                        print("Allowed Path has been set!")

                        messagebox.showinfo("info", "Allowed Path has been set!")

                        log_action('Allowed Path has been set!', item)

                    # starting

                    node_dict = {}

                    add_button = tk.Button(root, text='Add Node',command=lambda: add_node(tree, node_dict, None))
                    add_button.pack(side="left", padx=10, pady=5)

                    add_subnode_button = tk.Button(root, text='Add Subnode', command=lambda: add_subnode(tree, node_dict))
                    add_subnode_button.pack(side="left", padx=10, pady=5)

                    edit_button = tk.Button(root, text='Edit Node', command=lambda: edit_node(tree, node_dict))
                    edit_button.pack(side="left", padx=10, pady=5)

                    delete_button = tk.Button(root, text='Delete Node',   command=lambda: remove_node(tree, node_dict))
                    delete_button.pack(side="left", padx=10, pady=5)

                    detail_button = tk.Button(root, text='Show Details',   command=lambda: show_details(tree))  # , node_dict
                    detail_button.pack(side="left", padx=10, pady=5)

                    treeview_button = tk.Button(root, text='Set Treeview',     command=lambda: set_node_type(tree, node_dict, 'treeview'))
                    treeview_button.pack(side="left", padx=10, pady=5)

                    txt_button = tk.Button(root, text='Set TXT',command=lambda: set_node_type(tree, node_dict, node_type='txt'))
                    txt_button.pack(side="left", padx=10, pady=5)

                    save_button = tk.Button(root, text='Save Node-Treeview\n cltr + S ', command=lambda: save_treeview(tree, node_dict, file_path=None, override_existing=True))
                    save_button.pack(side="left", padx=10, pady=5)

                    load_button = tk.Button(root, text='Load Treeview', command=lambda: load_treeview(tree, node_dict, username_entry))
                    load_button.pack(side="left", padx=10, pady=5)

                    search_button = tk.Button(root, text='Search', command=lambda: search_node(tree, simpledialog.askstring("Input", "Enter search term:")))
                    search_button.pack(side="left", padx=10, pady=5)

                    # activ_tree_button = tk.Button(root, text='activ_tree', command=lambda: get_active_tree_name(trees, trees_names))
                    # activ_tree_button.pack(side="left", padx=10, pady=5)

                ##
                    # id_test_button = tk.Button(root, text='id_test', command=lambda: get_selected_id(tree))
                    # id_test_button.pack(side="left", padx=10, pady=5)

                    # focus_button = tk.Button(root, text='focus', command=lambda: find_and_focus(tree))
                    # focus_button.pack(side="left", padx=10, pady=5)

                    # button = tk.Button(root, text='Open Web Page', command=open_webpage)
                    # button.pack(side="left", padx=10, pady=5)

                    # history_button = tk.Button(root, text="History", command=show_file_content)

                    # history_button.place(x=10, y=10)

                    # set_path_button = tk.Button(root, text="Set Allowed Path", command=lambda: set_allowed_path(tree,))

                    # set_path_button.place(x=100, y=10)

                    # Assume 'treeview' is the Treeview widget
                    clear_button = tk.Button(root, text="Clear", command=lambda: clear_treeview(tree))
                    clear_button.pack(side="left", padx=10, pady=5)

                    # root.mainloop()
                    # return history_button, set_path_button

                    ##
                    # set_path_button.pack(side="left", padx=10, pady=5)
                    ##
                    # history_button.pack(side="left", padx=10, pady=5)

        ###########################################################################################
        # start: part of Login code:
        # set_path_button = tk.Button(root, text="Set Allowed Path", command=lambda: set_allowed_path(tree,))
        CSV_FILE = fr"C:\Users\Administrateur\OneDrive\Desktop\UMS_management\log_data.csv"  # to load the csv file

        def forgot_password2():
        
            # class belongs to opencv to be able to create a capturing video frame,
            cap = cv2.VideoCapture(0)
            # and (0) arg is the index of the cam (0 is the 1st cam , 1 is the 2nd cam ... )
        
            cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
            # the width and the height of the video capture frame or window
            cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
        
            counter = 0
            face_match = False
        
            # Read the photo database from CSV file
            logd= []
            with open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\log_data.csv', 'r') as file:
                # creating a csv reader or Cursor (read and extract data from each iterated row in "file")
                reader = csv.reader(file)
                for row in reader:
                    # append : add an ele to the end of "log_data = []" list
                    logd.append(row[4])
                    logd.append(row[5])
                print ("logD:",logd)
                print("path1:", logd[0])
        
            while True:  # infinite loop looping until a break statment comes
                # frame: read the video frame, ret: true if the frame was readed successfully
                ret, frame = cap.read()
        
                if ret:  # if ture
                    # every 30 (60 or 90) capture the facial recognition check is trigered
                    if counter % 30 == 0:
                        for photo_path in logd:
                            reference_img = cv2.imread(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\face_photo\reference.jpg')
                            #assigned any img in the list as a ref img to start with
        
                            try:
                                if DeepFace.verify(frame, reference_img.copy())['verified']:
        
                                    # The ['verified'] syntax is used to access the value associated with the 'verified' key in the result of the DeepFace.verify() function call.
        
                                    face_match = True
                                    this_photo_path=photo_path
                                    t =log_data[5]
        
                                    break
                                else:
                                    face_match = False
                            except:
                                face_match = False
        
                    counter += 1  # if false then pass to the next frame and after 30 captures the check proc trigered again
        
                    if face_match: 
                        
                        if logd[1]=="admin":  ##and row[4]==this_photo_path:
        
                            print("your face_photo is match !")
            
                            cv2.putText(frame, "Login successfully", (20, 450),           cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 255, 0), 3)
            
                            # delay=5000
            
                            root.destroy()
                            administrator_section(data)
                    
                        if logd[3]=="user": ## and row[4]==this_photo_path:
    
                            print("your face_photo is match !")
            
                            cv2.putText(frame, "Login successfully", (20, 450),           cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 255, 0), 3)
            
                            # delay=5000
            
                            root.destroy()
                            user_section()
                            
                            
                    else:
                        cv2.putText(frame, "Access Denied", (20, 450),       cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 0, 255), 3)
        
                    cv2.imshow("video", frame)  # the frame title is 'video'
        
                # exe (the break ) after 1ms from the key pressing ('q')
                key = cv2.waitKey(1)
                if key == ord("q"):  # if 'q'
                    break
        
            cv2.destroyAllWindows()  # close all opencv windows

        root = tk.Tk()

        user_type_var = tk.StringVar()

        combo = ttk.Combobox(root,  values=("admin", "user"))
        combo.grid(row=2, column=1, padx=10, pady=10, sticky="e")
        combo.current(0)

        def check_credentials(username_entry, password_entry, user_type_combobox, tree):
            username = username_entry.get()
            print("usen:", username)
            password = password_entry.get()
            print("pw:", password)

            # this methode is instead of combo.get()  , cause it didn't works

            def get_selected_value():
                selected_index = combo.current()  # Get the index of the selected item
                if selected_index >= 0:  # Make sure an item is selected
                    # Get the value using the index
                    selected_value = combo['values'][selected_index]
                    print("Selected value:", selected_value)
                else:
                    print("No item selected")

                return selected_value


           ## the login issue

            comv = get_selected_value()
            print("comv:", comv)
            # user_type =user_type_var.get()
            # print("type:",user_type_var)
            is_valid = False

            for item in tree.get_children():
                data = tree.item(item)["values"]

                ## here u should print for Debugging purpose
                

                if data[0] == username and data[1] == password and comv == data[5] == "user":
                    user_section()
                    break
                elif data[0] == username and data[1] == password and comv == data[5] == "admin":
                    administrator_section(data)
                    break
                else:
                    print("there is an error")

        def check_admin_credentials(admin_password_entry, tree):
            password = admin_password_entry.get()

            if password == "a":
                admin_login_window.destroy()
                logs_manager(tree)
            else:
                messagebox.showerror("Admin Login Failed",        "Incorrect password.")

        def open_admin_login_window(tree):
            global admin_login_window
            admin_login_window = tk.Toplevel(root)
            admin_login_window.title("Admin Login")

            admin_password_label = tk.Label(admin_login_window, text="Password:")
            admin_password_label.grid(row=0, column=0, padx=10, pady=10, sticky="e")
            admin_password_entry = tk.Entry(admin_login_window, show="*")
            admin_password_entry.grid(row=0, column=1, padx=10, pady=10)

            admin_login_button = tk.Button(admin_login_window, text="Login", command=lambda: check_admin_credentials(admin_password_entry, tree))
            admin_login_button.grid(row=1, column=0, columnspan=2, padx=10, pady=10)  # columnspan=2: this means that the button will appears inthe middle of the entry above

            admin_back_button = tk.Button(admin_login_window, text="Back", command=admin_login_window.destroy)
            admin_back_button.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

            admin_login_window.lift()

        def load_from_csv(tree):
            try:
                with open(CSV_FILE, "r", newline="") as f:  # 'r': open it in the read mode ; newline="" : to handel newline syntax in diff OS
                    reader = csv.reader(f)  # create a reader or a cursor
                    for row in reader:  # attach the cursor theads to each row in f
                        tree.insert("", "end", values=row)  # "": the 1st item will be in the root index position ,"end":the next item will be just at the side (means there is no espaces), "values=row": means, take the value from the current row (then put it into the current item)
            except FileNotFoundError:
                pass

        def save_data(entry_fields, tree):
            new_data = [entry.get() for entry in entry_fields]

            if "" in new_data:
                messagebox.showerror("Save Error", "All fields must be filled.")
                return

            tree.insert("", "end", values=new_data)

            # Save data to the CSV file
            with open(CSV_FILE, "a", newline="") as f:
                writer = csv.writer(f)
                # This is a method of the CSV writer object that writes a single row of data to the CSV file.
                writer.writerow(new_data)

            messagebox.showinfo("Save Successful", "Data saved successfully.")

        def update_csv(tree):
            data = [tree.item(item)["values"]
                              for item in tree.get_children()]  # issue

            with open(CSV_FILE, "w", newline="") as f:
                writer = csv.writer(f)
                # This is a method of the CSV writer object that writes multi rows of data to the CSV file.
                writer.writerows(data)

        def edit_data(entry_fields, tree):
            selected_item = tree.selection()
            if not selected_item:
                messagebox.showerror("Edit Error", "No item is selected.")
                return

            new_data = [entry.get() for entry in entry_fields]  # issue

            if "" in new_data:
                messagebox.showerror("Edit Error", "All fields must be filled.")
                return

            tree.item(selected_item, values=new_data)

            # Update data in the CSV file
            update_csv(tree)
            messagebox.showinfo("Edit Successful", "Data edited successfully.")

        def delete_data(tree):
            selected_item = tree.selection()
            if not selected_item:
                messagebox.showerror("Delete Error", "No item is selected.")
                return

            tree.delete(selected_item)

            # Update data in the CSV file
            update_csv(tree)
            messagebox.showinfo("Delete Successful",   "Data deleted successfully.")

        def reset_data(entry_fields):
            for entry in entry_fields:
                entry.delete(0, tk.END)

        def logs_manager(tree):
            logs_window = tk.Toplevel(root)
            logs_window.title("Logs Manager")

            # Entry fields
            field_names = ["Name/Username", "Password", "Phone", "CIN", "img_path", "user_type", "allowed_path"]
            entry_fields = []

            for index, field_name in enumerate(field_names):
                label = tk.Label(logs_window, text=f"{field_name}:")
                label.grid(row=index, column=0, padx=10, pady=10, sticky="e")
                entry = tk.Entry(logs_window)
                entry.grid(row=index, column=1, padx=10, pady=10)
                entry_fields.append(entry)

            # Treeview
            logs_tree = ttk.Treeview(logs_window, columns=field_names, show="headings")
            for name in field_names + ["img_path"]:  # + ["img_path"]:    cause it's require file dialog  and it's not a simple entry type
                logs_tree.heading(name, text=name)
            logs_tree.grid(row=10, column=0, columnspan=4,
                           padx=10, pady=10, sticky="news")

            # Load data from CSV into Treeview
            for item in tree.get_children():
                logs_tree.insert("", "end", values=tree.item(item)["values"])

            # Upload Image function
            def upload_image():
                selected_item = logs_tree.focus()
                if selected_item:
                    image_path = filedialog.askopenfilename(title="Select an Image", filetypes=[
                                                            ("Image Files", "*.png *.jpg *.jpeg")])
                    if image_path:
                        # Update the img_path value in the selected item
                        logs_tree.set(selected_item, "img_path", image_path)

                        # Update the img_path value in the CSV file
                        update_csv(logs_tree)

            # Buttons
            save_button = tk.Button(logs_window, text="Save", command=lambda: save_data(entry_fields, logs_tree))
            save_button.grid(row=15, column=0, padx=10, pady=10, sticky="news")

            edit_button = tk.Button(logs_window, text="Edit", command=lambda: edit_data(entry_fields, logs_tree))
            edit_button.grid(row=15, column=1, padx=10, pady=10, sticky="news")

            delete_button = tk.Button(logs_window, text="Delete", command=lambda: delete_data(logs_tree))
            delete_button.grid(row=15, column=2, padx=10,  pady=10, sticky="news")

            reset_button = tk.Button(logs_window, text="Reset", command=lambda: reset_data(entry_fields))
            reset_button.grid(row=15, column=3, padx=10, pady=10, sticky="news")

            # Upload Image button
            upload_button = tk.Button(logs_window, text="Upload Image", command=upload_image)
            upload_button.grid(row=16, column=0, columnspan=4,  padx=10, pady=10, sticky="news")

        root.title("Grades Section Login")
        username_label = tk.Label(root, text="Username:")
        username_label.grid(row=0, column=0, padx=10, pady=10, sticky="e")
        username_entry = tk.Entry(root)
        username_entry.grid(row=0, column=1, padx=10, pady=10)
        password_label = tk.Label(root, text="Password:")
        password_label.grid(row=1, column=0, padx=10, pady=10, sticky="e")
        password_entry = tk.Entry(root, show="*")
        password_entry.grid(row=1, column=1, padx=10, pady=10)

        # Initialize tree variable and load data from CSV
        tree = ttk.Treeview()
        load_from_csv(tree)

        data = None
        # data access (start)
        for item in tree.get_children():
            data = tree.item(item)["values"]

            print('data inLood', data)

        print("data outLoop:", data)
        # data access (end)


        login_button = tk.Button(root, text="Login", command=lambda: check_credentials(username_entry, password_entry, combo, tree))
        login_button.grid(row=3, column=0, columnspan=2, padx=10, pady=10)
        ##forgot_password_button = tk.Button(##    root, text="Forgot Password?", command=forgot_password2)
        ##forgot_password_button.grid(##    row=4, column=0, columnspan=2, padx=10, pady=10)
        admin_login_button = tk.Button(root, text="Admin Login", command=lambda: open_admin_login_window(tree))
        admin_login_button.grid(row=5, column=0, columnspan=2, padx=10, pady=10)
        root.mainloop()


    # end:  part of Logs Manager code


    # issues: hot key ctr + s to save the change when building a node-treeview

    ##action7 just for testing
    def action7():

        webbrowser.open('http://localhost/Libraryms-PHP-copy3/')


# START (after the main)

    # Create a new Tkinter window
    root = tk.Tk()
    root.title('System Dashbord')
    # Configure the rows and columns to expand with the window


    # Create the top frame and the title label
    topframe = tk.Frame(root)
    topframe.grid(row=0, column=0, columnspan=3, sticky='ew')
    title = tk.Label(topframe, text="University Management System", font=("Helvetica", 16))
    title.pack(pady=10)


    # Create a bottom frame and the "visit this site" label
    bottomframe = tk.Frame(root)
    bottomframe.grid(row=7, column=0, columnspan=3, sticky='ew')
    visit_label = tk.Label(bottomframe, text="To Access Grades SYS Remotely Plz Visit This Web Site:\n ", font=("Arial", 10, "bold"), fg="green")
    
    # Create a separate label for the red URL
    url_label = tk.Label(bottomframe, text="http://127.0.0.1:5000/", font=("Arial", 10, "bold"), fg="red")
    
    visit_label.pack(pady=(30, 0))
    url_label.pack(pady=(0, 30))


    ## image starts 


    from tkinter import Tk, Label, PhotoImage
    from PIL import Image, ImageTk

    # Load the images and resize them
    # image_size = (100, 100)  # Set the desired image size

    # Load the images
    image1 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\1_aboutUni_icon.png'))





    import os

    # Get the directory where your script (v28_MM_fix.py) is located
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    # Build the path to the icon dynamically
    icon_path = os.path.join(BASE_DIR, "icons_Gp5", "2_student_icon.png")

    # Now open it
    try:
        image2 = ImageTk.PhotoImage(Image.open(icon_path))
    except FileNotFoundError:
        print(f"Error: Could not find the icon at {icon_path}")
        # Fallback: create a blank image or skip    image3 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\3_staff_icon.png'))
    
    
    
    
    
    
    
    
    
    # image4 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\4_uniRes_icon.gif'))
    # image5 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\5_timetable_icon.png'))
    image6 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\6_grades_icon.png'))
    # image7 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\7_lib_icon.png'))
    # Create the buttons, assigning each one its action

    # Function to load and resize images
    # def load_and_resize_image(image_path, width, height):
    # image = Image.open(image_path)
    # image = image.resize((width, height), Image.ANTIALIAS)
    # return ImageTk.PhotoImage(image)
    ##
    # Image dimensions for buttons
    # image_width, image_height = 100, 100
    ##
    # Load and resize the images
    # image2 = load_and_resize_image(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\2_student_icon.png', image_width, image_height)
    # image3 = load_and_resize_image(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\3_staff_icon.png', image_width, image_height)
    # image6 = load_and_resize_image(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\6_grades_icon.png', image_width, image_height)
    ##
    # ...
    ##


# Load the GIF using PIL.ImageTk.PhotoImage

    # gif_path1 = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\4_uniRes_icon.gif'
    # gif_frames1 = Image.open(gif_path1)
    # gif_frames1 = [ImageTk.PhotoImage(frame) for frame in ImageSequence.Iterator(gif_frames1)]

    gif_path4 = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\4_uniRes_icon.gif'
    gif_frames4 = Image.open(gif_path4)
    gif_frames4 = [ImageTk.PhotoImage(frame) for frame in ImageSequence.Iterator(gif_frames4)]

    gif_path5 = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\5_icon_timetable.gif'
    gif_frames5 = Image.open(gif_path5)
    gif_frames5 = [ImageTk.PhotoImage(frame) for frame in ImageSequence.Iterator(gif_frames5)]

    gif_path7 = fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\7_lib_icon.gif'
    gif_frames7 = Image.open(gif_path7)
    gif_frames7 = [ImageTk.PhotoImage(frame) for frame in ImageSequence.Iterator(gif_frames7)]

    # Create the buttons, assigning each one its action
    style = ttk.Style()
    style.configure('Hover.TButton', background='gray90', padding=20)


    button1 = ttk.Button(root,  command=action1, image=image1,style='Hover.TButton', cursor='hand2')
    button1.grid(row=1, column=0, sticky='nsew')
    label1 = tk.Label(root, text="University Details")
    label1.grid(row=2, column=0, sticky='nsew')

    button2 = ttk.Button(root,  command=action2, image=image2,style='Hover.TButton', cursor='hand2')
    button2.grid(row=1, column=1, sticky='nsew')
    label2 = tk.Label(root, text="Student Management", )
    label2.grid(row=2, column=1, sticky='nsew')

    button3 = ttk.Button(root,  command=action3, image=image1,style='Hover.TButton', cursor='hand2')
    button3.grid(row=1, column=2, sticky='nsew')
    label3 = tk.Label(root, text="Staff Management")
    label3.grid(row=2, column=2, sticky='nsew')

    button4 = ttk.Button(root,  command=action4,style='Hover.TButton', cursor='hand2')
    button4.grid(row=3, column=0, sticky='nsew')
    label4 = tk.Label(root, text="University Resources Management")
    label4.grid(row=4, column=0, sticky='nsew')

    button5 = ttk.Button(root,  command=action5,style='Hover.TButton', cursor='hand2')
    button5.grid(row=3, column=1, sticky='nsew')
    label5 = tk.Label(root, text="Timetables Management")
    label5.grid(row=4, column=1, sticky='nsew')

    button6 = ttk.Button(root,  command=action6, image=image6,style='Hover.TButton', cursor='hand2')
    button6.grid(row=3, column=2, sticky='nsew')
    label6 = tk.Label(root, text="Student Grades Management")
    label6.grid(row=4, column=2, sticky='nsew')

    button7 = ttk.Button(root,  command=action7,style='Hover.TButton', cursor='hand2')
    button7.grid(row=5, column=1, sticky='nsew')
    label7 = tk.Label(root, text="Library")
    label7.grid(row=6, column=1, sticky='nsew')

    # Function to update the button image with the next frame in the GIF


    def update_gif(button, frames, index):
        button.config(image=frames[index])
        index = (index + 1) % len(frames)
        root.after(100, update_gif, button, frames, index)

    # Start the GIF animation for button1 and button2
    # update_gif(button1, gif_frames1, 0)
    update_gif(button4, gif_frames4, 0)
    update_gif(button5, gif_frames5, 0)
    update_gif(button7, gif_frames7, 0)

    # Start the Tkinter event loop
    root.mainloop()



    ##import tkinter as tk
    ##from tkinter import ttk
    ##from PIL import Image, ImageTk, ImageSequence
    ##
    ### Define your action functions here (e.g., action1, action2, etc.)
    ##
    ### Create a new Tkinter window
    ##root = tk.Tk()
    ##root.title('System Dashboard')
    ##
    ### Configure the rows and columns to expand with the window
    ##for i in range(7):
    ##    root.grid_rowconfigure(i, weight=1)
    ##for i in range(3):
    ##    root.grid_columnconfigure(i, weight=1)
    ##
    ### Create the top frame and the title label
    ##topframe = tk.Frame(root)
    ##topframe.grid(row=0, column=0, columnspan=3, sticky='ew')
    ##title = tk.Label(topframe, text="University Management System", font=("Helvetica", 16))
    ##title.pack(pady=10)
    ##
    ### Create a bottom frame and the "visit this site" label
    ##bottomframe = tk.Frame(root)
    ##bottomframe.grid(row=7, column=0, columnspan=3, sticky='ew')
    ##visit_label = tk.Label(bottomframe, text="To Access Grades SYS Remotely Plz Visit This Web Site:\n ", font=("Arial", 10, "bold"), fg="green")
    ##
    ### Create a separate label for the red URL
    ##url_label = tk.Label(bottomframe, text="http://127.0.0.1:5000/", font=("Arial", 10, "bold"), fg="red")
    ##
    ##visit_label.pack(pady=(30, 0))
    ##url_label.pack(pady=(0, 30))
    ##
    ### Load the images and resize them
    ##image1 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\1_aboutUni_icon.png'))
    ##image2 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\2_student_icon.png'))
    ##image3 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\3_staff_icon.png'))
    ##image4 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\4_uniRes_icon.gif'))
    ##image5 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\5_icon_timetable.gif'))
    ##image6 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\6_grades_icon.png'))
    ##image7 = ImageTk.PhotoImage(Image.open(fr'C:\Users\Administrateur\OneDrive\Desktop\UMS_management\icons_Gp5\icons_Gp55\7_lib_icon.gif'))
    ##
    ### Create the buttons, assigning each one its action
    ##button1 = ttk.Button(root, command=action1, image=image1, style='Hover.TButton', cursor='hand2')
    ##button1.grid(row=1, column=0, sticky='nsew')
    ##label1 = tk.Label(root, text="University Details")
    ##label1.grid(row=2, column=0, sticky='nsew')
    ##
    ##button2 = ttk.Button(root, command=action2, image=image2, style='Hover.TButton', cursor='hand2')
    ##button2.grid(row=1, column=1, sticky='nsew')
    ##label2 = tk.Label(root, text="Student Management")
    ##label2.grid(row=2, column=1, sticky='nsew')
    ##
    ##button3 = ttk.Button(root, command=action3, image=image3, style='Hover.TButton', cursor='hand2')
    ##button3.grid(row=1, column=2, sticky='nsew')
    ##label3 = tk.Label(root, text="Staff Management")
    ##label3.grid(row=2, column=2, sticky='nsew')
    ##
    ##button4 = ttk.Button(root, command=action4, image=image4, style='Hover.TButton', cursor='hand2')
    ##button4.grid(row=3, column=0, sticky='nsew')
    ##label4 = tk.Label(root, text="University Resources Management")
    ##label4.grid(row=4, column=0, sticky='nsew')
    ##
    ##button5 = ttk.Button(root, command=action5, image=image5, style='Hover.TButton', cursor='hand2')
    ##button5.grid(row=3, column=1, sticky='nsew')
    ##label5 = tk.Label(root, text="Timetables Management")
    ##label5.grid(row=4, column=1, sticky='nsew')
    ##
    ##button6 = ttk.Button(root, command=action6, image=image6, style='Hover.TButton', cursor='hand2')
    ##button6.grid(row=3, column=2, sticky='nsew')
    ##label6 = tk.Label(root, text="Student Grades Management")
    ##label6.grid(row=4, column=2, sticky='nsew')
    ##
    ##button7 = ttk.Button(root, command=action7, image=image7, style='Hover.TButton', cursor='hand2')
    ##button7.grid(row=5, column=1, sticky='nsew')
    ##label7 = tk.Label(root, text="Library")
    ##label7.grid(row=6, column=1, sticky='nsew')
    ##
    ### Function to update the button image with the next frame in the GIF
    ##def update_gif(button, frames, index):
    ##    button.config(image=frames[index])
    ##    index = (index + 1) % len(frames)
    ##    root.after(100, update_gif, button, frames, index)
    ##
    ### Start the GIF animation for button4, button5, and button7
    ##update_gif(button4, gif_frames4, 0)
    ##update_gif(button5, gif_frames5, 0)
    ##update_gif(button7, gif_frames7, 0)
    ##
    ### Start the Tkinter event loop
    ##root.mainloop()



# Here put the source code of main1  (end)


# cut(start)


def main():
    global root

    root = tk.Tk()
    root.title("Login Window")

    username_label = tk.Label(root, text="Username:")
    username_label.grid(row=0, column=0, padx=10, pady=10, sticky="e")
    username_entry = tk.Entry(root)
    username_entry.grid(row=0, column=1, padx=10, pady=10)

    password_label = tk.Label(root, text="Password:")
    password_label.grid(row=1, column=0, padx=10, pady=10, sticky="e")
    password_entry = tk.Entry(root, show="*")
    password_entry.grid(row=1, column=1, padx=10, pady=10)

    # Initialize tree variable and load data from CSV
    tree = ttk.Treeview()
    load_from_csv(tree)

    login_button = tk.Button(root, text="Login", command=lambda: check_credentials(username_entry, password_entry, tree))
    login_button.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

    forgot_password_button = tk.Button(root, text="Forgot Password?", command=forgot_password)
    forgot_password_button.grid(row=3, column=0, columnspan=2, padx=10, pady=10)

    admin_login_button = tk.Button(root, text="Admin Login", command=lambda: open_admin_login_window(tree))
    admin_login_button.grid(row=4, column=0, columnspan=2, padx=10, pady=10)

    root.mainloop()


if __name__ == "__main__":
    main()


# it DataBase (log_data.csv)
